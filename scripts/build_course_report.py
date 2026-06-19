from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "assets"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
OUT = REPORT_DIR / "智能数据守护者系统课程设计报告.docx"


def font(size=24, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in candidates:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def draw_box(draw, box, text, fill="#ffffff", outline="#2563eb", text_fill="#111827", width=2, radius=8, fnt=None):
    fnt = fnt or font(22)
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    lines = wrap_text(text, fnt, box[2] - box[0] - 24)
    line_h = fnt.size + 7
    total_h = len(lines) * line_h
    y = box[1] + ((box[3] - box[1] - total_h) / 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = box[0] + ((box[2] - box[0] - (bbox[2] - bbox[0])) / 2)
        draw.text((x, y), line, font=fnt, fill=text_fill)
        y += line_h


def wrap_text(text, fnt, max_width):
    lines = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            if fnt.getlength(trial) <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines or [""]


def arrow(draw, start, end, fill="#334155", width=3):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 12
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        x = x2 + length * math.cos(angle + delta)
        y = y2 + length * math.sin(angle + delta)
        draw.line([(x2, y2), (x, y)], fill=fill, width=width)


def save_diagram(name, title, boxes, arrows, size=(1600, 1000)):
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", size, "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.text((50, 32), title, font=font(36, True), fill="#0f172a")
    for item in boxes:
        draw_box(draw, item["box"], item["text"], item.get("fill", "#ffffff"), item.get("outline", "#2563eb"), fnt=font(item.get("font", 22), item.get("bold", False)))
    for item in arrows:
        arrow(draw, item["start"], item["end"], item.get("fill", "#334155"))
    path = DIAGRAM_DIR / name
    img.save(path)
    return path


def build_diagrams():
    use_case_boxes = [
        {"box": (60, 170, 230, 260), "text": "普通员工", "fill": "#e0f2fe", "outline": "#0284c7", "bold": True},
        {"box": (60, 380, 230, 470), "text": "部门主管", "fill": "#ecfdf5", "outline": "#059669", "bold": True},
        {"box": (60, 590, 230, 680), "text": "数据安全员", "fill": "#fef9c3", "outline": "#ca8a04", "bold": True},
        {"box": (60, 800, 230, 890), "text": "系统管理员", "fill": "#fee2e2", "outline": "#dc2626", "bold": True},
        {"box": (420, 120, 710, 205), "text": "文本检测 / 文件检测"},
        {"box": (760, 120, 1050, 205), "text": "查看检测结果与历史"},
        {"box": (420, 320, 710, 405), "text": "审核申请"},
        {"box": (760, 320, 1050, 405), "text": "风险处理 / 审核记录"},
        {"box": (420, 520, 710, 605), "text": "AI审核"},
        {"box": (760, 520, 1050, 605), "text": "敏感词库 / 风险等级"},
        {"box": (420, 740, 710, 825), "text": "用户管理 / 角色管理"},
        {"box": (760, 740, 1050, 825), "text": "权限管理 / 日志 / 配置"},
    ]
    use_case_arrows = [
        {"start": (230, 215), "end": (420, 160)}, {"start": (230, 215), "end": (760, 160)},
        {"start": (230, 425), "end": (420, 360)}, {"start": (230, 425), "end": (760, 360)},
        {"start": (230, 635), "end": (420, 560)}, {"start": (230, 635), "end": (760, 560)},
        {"start": (230, 845), "end": (420, 780)}, {"start": (230, 845), "end": (760, 780)},
    ]
    save_diagram("uml_use_case.png", "图 1 系统 UML 用例图", use_case_boxes, use_case_arrows)

    class_boxes = [
        {"box": (80, 170, 360, 315), "text": "Controller层\nAuthController\nDetectionController\nAuditController\nAdminController", "fill": "#eff6ff"},
        {"box": (480, 170, 760, 315), "text": "Service层\nAuthService\nDetectionService\nAiReviewService\nAuditService", "fill": "#ecfdf5"},
        {"box": (880, 170, 1160, 315), "text": "Mapper层\nSysUserMapper\nDetectTaskMapper\nAuditTaskMapper\nOperationLogMapper", "fill": "#fff7ed"},
        {"box": (1280, 170, 1530, 315), "text": "数据库\nMySQL 8.0\nRBAC表\n检测审核表\n日志表", "fill": "#fef2f2"},
        {"box": (480, 500, 760, 650), "text": "核心实体\nSysUser\nDetectTask\nDetectResult\nAuditTask\nWarningRecord", "fill": "#f8fafc"},
        {"box": (880, 500, 1160, 650), "text": "外部服务\nDeepSeek API\n文件解析\nJWT/BCrypt", "fill": "#f0f9ff"},
    ]
    class_arrows = [
        {"start": (360, 240), "end": (480, 240)}, {"start": (760, 240), "end": (880, 240)},
        {"start": (1160, 240), "end": (1280, 240)}, {"start": (620, 315), "end": (620, 500)},
        {"start": (760, 575), "end": (880, 575)},
    ]
    save_diagram("uml_class.png", "图 2 系统分层类图", class_boxes, class_arrows)

    flow_boxes = [
        {"box": (80, 170, 330, 255), "text": "员工提交文本/文件"},
        {"box": (420, 170, 670, 255), "text": "规则检测\n手机号/身份证/邮箱/银行卡/敏感词"},
        {"box": (760, 170, 1010, 255), "text": "DeepSeek AI审核\n生成建议"},
        {"box": (1100, 170, 1350, 255), "text": "生成风险等级\n检测结果入库"},
        {"box": (420, 420, 670, 505), "text": "高风险触发预警"},
        {"box": (760, 420, 1010, 505), "text": "主管审核\n通过/驳回"},
        {"box": (1100, 420, 1350, 505), "text": "员工查看审核结果\n管理员查看日志"},
    ]
    flow_arrows = [
        {"start": (330, 212), "end": (420, 212)}, {"start": (670, 212), "end": (760, 212)},
        {"start": (1010, 212), "end": (1100, 212)}, {"start": (1225, 255), "end": (545, 420)},
        {"start": (670, 462), "end": (760, 462)}, {"start": (1010, 462), "end": (1100, 462)},
    ]
    save_diagram("workflow.png", "图 3 业务流程图", flow_boxes, flow_arrows)

    er_boxes = [
        {"box": (80, 130, 330, 270), "text": "sys_user\n用户\nrole_code\ndepartment_id"},
        {"box": (430, 130, 680, 270), "text": "sys_role\n角色\nrole_code\nrole_name"},
        {"box": (780, 130, 1030, 270), "text": "sys_permission\n权限\npermission_code\nmenu_path"},
        {"box": (1130, 130, 1450, 270), "text": "sys_department\n部门\nparent_id"},
        {"box": (80, 430, 330, 590), "text": "detect_task\n检测任务\nuser_id\ncontent_type\nstatus"},
        {"box": (430, 430, 680, 590), "text": "detect_result\n检测结果\ntask_id\nrisk_level\nrisk_score"},
        {"box": (780, 430, 1030, 590), "text": "audit_task\n审核任务\ndetect_task_id\nreviewer_id"},
        {"box": (1130, 430, 1450, 590), "text": "warning_record\n预警记录\ndetect_task_id\nwarning_level"},
        {"box": (80, 740, 330, 900), "text": "sensitive_word\n敏感词\ncategory_id\nrisk_level"},
        {"box": (430, 740, 680, 900), "text": "sensitive_category\n敏感词分类"},
        {"box": (780, 740, 1030, 900), "text": "operation_log\n操作日志\nuser_id\noperation"},
        {"box": (1130, 740, 1450, 900), "text": "risk_level_config\n风险等级配置"},
    ]
    er_arrows = [
        {"start": (330, 200), "end": (430, 200)}, {"start": (680, 200), "end": (780, 200)},
        {"start": (330, 500), "end": (430, 500)}, {"start": (330, 500), "end": (780, 500)},
        {"start": (330, 500), "end": (1130, 500)}, {"start": (330, 200), "end": (205, 430)},
        {"start": (330, 820), "end": (430, 820)}, {"start": (330, 200), "end": (780, 820)},
        {"start": (1130, 200), "end": (330, 200)},
    ]
    save_diagram("er_diagram.png", "图 4 数据库 E-R 图", er_boxes, er_arrows, size=(1600, 1000))


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text))
    doc.add_paragraph()
    return table


def set_doc_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(15, 23, 42)


def paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(21)
    r = p.add_run("· " + text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)


def add_image(doc, path, caption, width=6.2):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(caption).bold = True
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实践课程设计（专业领域）报告")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(26)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("智能数据守护者系统")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)
    doc.add_paragraph()
    rows = [
        ("课程名称", "实践课程设计（专业领域）"),
        ("项目名称", "智能数据守护者系统"),
        ("技术栈", "Spring Boot 3 / MyBatis Plus / Vue3 / MySQL / DeepSeek API"),
        ("班级", "计算机23级（请填写）"),
        ("学号", "请填写"),
        ("姓名", "请填写"),
        ("指导教师", "请填写"),
        ("完成日期", "2026年6月"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], row[0], True)
        set_cell_text(table.rows[i].cells[1], row[1])
    doc.add_page_break()


def build_report():
    REPORT_DIR.mkdir(exist_ok=True)
    build_diagrams()
    doc = Document()
    set_doc_styles(doc)
    cover(doc)

    doc.add_heading("目录", level=1)
    for item in [
        "1 项目概述",
        "2 需求分析",
        "3 系统总体设计",
        "4 数据库设计",
        "5 系统详细设计与实现",
        "6 系统运行样例",
        "7 测试与结果分析",
        "8 总结",
    ]:
        paragraph(doc, item)
    doc.add_page_break()

    doc.add_heading("1 项目概述", level=1)
    paragraph(doc, "智能数据守护者系统面向企业日常办公数据安全管理场景，解决内部通信、资料共享和外部发布过程中敏感信息泄露风险。系统采用前后端分离架构，后端基于 Spring Boot 3、MyBatis Plus、MySQL、JWT 与 BCrypt 实现业务与权限控制，前端基于 Vue3、Vite、Element Plus 与 ECharts 实现多角色操作界面，并接入 DeepSeek API 完成 AI 辅助审核。")
    paragraph(doc, "系统围绕普通员工、部门主管、数据安全员、系统管理员四类角色展开，覆盖文本检测、文件检测、AI 审核、主管审核、风险预警、敏感词维护、权限管理、日志审计等功能。")

    doc.add_heading("2 需求分析", level=1)
    doc.add_heading("2.1 角色与权限", level=2)
    add_table(doc, ["角色", "主要职责", "权限范围"], [
        ["普通员工", "提交检测内容", "文本检测、文件检测、查看检测结果、查看历史记录、个人中心"],
        ["部门主管", "审核员工内容", "审核申请、部门统计、风险处理、审核记录"],
        ["数据安全员", "企业数据安全管理", "AI审核、敏感词库管理、风险预警管理、风险等级配置"],
        ["系统管理员", "系统运维", "用户管理、角色管理、权限管理、日志管理、系统配置"],
    ])
    doc.add_heading("2.2 功能需求", level=2)
    for text in [
        "用户认证模块：实现登录、注册、验证码、JWT 鉴权、密码加密、个人资料维护。",
        "敏感数据检测模块：支持文本和 txt、pdf、doc、docx 文件检测，识别手机号、身份证号、邮箱、银行卡、地址和敏感词。",
        "AI 审核模块：接入 DeepSeek API，根据审核场景、资料类型、发布范围输出结构化审核建议。",
        "审核管理模块：部门主管查看检测详情、原文、脱敏内容、风险命中和 AI 建议，并执行通过或驳回。",
        "风险预警模块：高风险和严重风险检测任务自动生成预警记录。",
        "RBAC 权限模块：按角色控制前端菜单、路由和后端接口访问。",
        "系统日志模块：记录登录、检测、审核、预警处理、敏感词维护和管理员配置等关键操作。",
    ]:
        bullet(doc, text)

    doc.add_heading("3 系统总体设计", level=1)
    paragraph(doc, "系统采用浏览器/服务器模式，前端 Vue3 通过 Axios 调用后端 REST API，后端通过 MyBatis Plus 访问 MySQL 数据库。敏感信息识别由规则检测与敏感词库共同完成，AI 审核通过 DeepSeek API 输出审核建议。")
    add_image(doc, DIAGRAM_DIR / "uml_use_case.png", "图 1 系统 UML 用例图")
    add_image(doc, DIAGRAM_DIR / "uml_class.png", "图 2 系统分层类图")
    add_image(doc, DIAGRAM_DIR / "workflow.png", "图 3 系统业务流程图")

    doc.add_heading("4 数据库设计", level=1)
    paragraph(doc, "数据库采用 MySQL 8.0，围绕 RBAC 权限、检测任务、审核任务、风险预警、敏感词库、系统配置和操作日志建立数据表。核心表包括 sys_user、sys_role、sys_permission、detect_task、detect_result、audit_task、audit_record、warning_record、sensitive_word、operation_log 等。")
    add_image(doc, DIAGRAM_DIR / "er_diagram.png", "图 4 数据库 E-R 图")
    add_table(doc, ["数据表", "说明"], [
        ["sys_user", "系统用户表，存储账号、密码、部门、角色、状态"],
        ["detect_task", "检测任务表，记录员工提交的文本或文件内容"],
        ["detect_result", "检测结果表，记录风险等级、评分、命中摘要、AI 建议"],
        ["audit_task / audit_record", "主管审核任务和审核记录"],
        ["warning_record", "高风险和严重风险预警记录"],
        ["sensitive_word / sensitive_category", "敏感词及分类维护"],
        ["operation_log", "系统审计日志"],
        ["risk_level_config", "风险等级配置"],
    ])

    doc.add_heading("5 系统详细设计与实现", level=1)
    doc.add_heading("5.1 后端实现", level=2)
    paragraph(doc, "后端按照 Controller、Service、Mapper、Entity 分层实现。Controller 负责接收前端请求，Service 负责业务流程编排，Mapper 负责数据库访问，Entity 与数据库表对应。系统通过 Spring Security 和 JWT 实现登录态校验，通过 @PreAuthorize 对不同角色接口进行权限控制。")
    doc.add_heading("5.2 前端实现", level=2)
    paragraph(doc, "前端根据登录用户 roleCode 动态生成菜单和路由。普通员工、部门主管、数据安全员、系统管理员分别进入不同首页，并只能访问各自功能。Element Plus 用于表单、表格、弹窗、上传控件和标签展示，ECharts 用于统计分析展示。")
    doc.add_heading("5.3 敏感检测与 AI 审核", level=2)
    paragraph(doc, "规则检测通过正则表达式识别手机号、身份证号、邮箱、银行卡号、地址等结构化敏感信息，并结合数据库中的敏感词库进行命中统计。DeepSeek AI 审核在规则检测基础上结合业务场景输出风险等级、是否建议发布、风险依据、修改建议和审核结论。")

    doc.add_heading("6 系统运行样例", level=1)
    screenshots = [
        ("01-login.png", "图 5 登录与验证码页面"),
        ("02-employee-detect.png", "图 6 普通员工文本/文件检测页面"),
        ("03-employee-history.png", "图 7 普通员工检测结果与历史页面"),
        ("04-employee-profile.png", "图 8 普通员工个人中心页面"),
        ("05-manager-audit.png", "图 9 部门主管审核申请页面"),
        ("06-manager-stats.png", "图 10 部门主管部门统计页面"),
        ("07-manager-risk.png", "图 11 部门主管风险处理页面"),
        ("08-manager-records.png", "图 12 部门主管审核记录页面"),
        ("09-security-ai.png", "图 13 数据安全员 AI 审核页面"),
        ("10-security-sensitive.png", "图 14 数据安全员敏感词库管理页面"),
        ("11-security-warnings.png", "图 15 数据安全员风险预警管理页面"),
        ("12-security-risk-levels.png", "图 16 数据安全员风险等级配置页面"),
        ("13-admin-users.png", "图 17 系统管理员用户管理页面"),
        ("14-admin-roles.png", "图 18 系统管理员角色管理页面"),
        ("15-admin-permissions.png", "图 19 系统管理员权限管理页面"),
        ("16-admin-logs.png", "图 20 系统管理员日志管理页面"),
        ("17-admin-config.png", "图 21 系统管理员系统配置页面"),
    ]
    for file, cap in screenshots:
        add_image(doc, SCREENSHOT_DIR / file, cap, width=6.3)

    doc.add_heading("7 测试与结果分析", level=1)
    add_table(doc, ["测试项", "测试账号", "输入/操作", "预期结果"], [
        ["文本检测", "employee", "输入包含手机号、客户名单、财务数据的文本", "生成检测结果、风险评分、AI 建议和审核任务"],
        ["文件检测", "employee", "上传 txt/pdf/doc/docx 文件", "提取文本并完成敏感检测"],
        ["主管审核", "manager", "查看审核详情并通过/驳回", "员工历史显示审核状态和意见"],
        ["风险预警", "manager/security", "检测高风险或严重风险内容", "生成预警记录并可处理"],
        ["AI审核", "security", "输入文本或上传文件", "DeepSeek 返回结构化审核结果"],
        ["敏感词管理", "security", "新增、修改、删除敏感词", "敏感词库更新并影响检测命中"],
        ["权限管理", "admin", "维护权限编码与菜单路径", "用于展示 RBAC 功能点字典"],
        ["日志管理", "admin", "查看操作日志", "记录登录、检测、审核、配置等关键操作"],
    ])
    paragraph(doc, "测试结果表明，系统能够完成从员工提交检测、系统自动识别、AI 审核建议、风险预警、主管审核、员工查看结果、管理员审计日志的完整闭环。")

    doc.add_heading("8 总结", level=1)
    paragraph(doc, "本项目实现了一个面向企业办公数据安全场景的智能数据守护者系统。系统综合使用 Spring Boot、MyBatis Plus、Vue3、MySQL、JWT、DeepSeek API 等技术，完成了敏感信息识别、多角色协同审核、AI 辅助判断、风险预警和日志审计等功能。通过本次课程设计，进一步掌握了前后端分离开发、RBAC 权限控制、数据库建模、文件解析、AI 接口集成和软件工程文档编写方法。")
    paragraph(doc, "后续可继续扩展 OCR 图片识别、多模型接入、自动生成安全报告、角色权限动态分配和更细粒度的数据脱敏策略。")

    doc.save(OUT)


if __name__ == "__main__":
    build_report()
    print(OUT)
