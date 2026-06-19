from __future__ import annotations

from pathlib import Path
import math
import re

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

import build_requested_template_report as base


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
SCREENSHOT_DIR = REPORT_DIR / "assets" / "screenshots"
DIAGRAM_DIR = REPORT_DIR / "assets" / "standard-black-diagrams"
OUT = REPORT_DIR / "智能数据守护者系统课程设计报告-最终无框流程图版.docx"


FEATURES = [
    ("登录注册模块", [
        ("验证码获取", "未登录用户", "用户进入登录页后，系统先生成验证码图片和验证码标识。验证码用于约束登录请求必须来自真实页面交互，减少脚本反复尝试密码的风险。该功能不写入业务表，但会参与登录校验。", "LoginView.vue", "/auth/captcha", "CaptchaService", ["验证码缓存"], "访问登录页", "返回验证码图片"),
        ("用户登录", "系统用户", "登录功能完成用户身份确认。用户提交用户名、密码和验证码后，后端校验验证码、账号状态和密码，校验通过后签发JWT令牌并返回角色信息。前端根据角色跳转到普通员工、部门主管、数据安全员或系统管理员首页。", "LoginView.vue", "/auth/login", "AuthService、JwtTokenProvider", ["sys_user", "operation_log"], "用户名、密码、验证码", "JWT令牌、角色信息"),
        ("用户注册", "新用户", "注册功能用于创建普通业务账号。后端需要检查用户名是否重复，密码必须加密保存。课程演示账号由系统初始化自动创建，注册功能体现系统对新增用户的扩展能力。", "LoginView.vue", "/auth/register", "AuthService", ["sys_user"], "注册表单", "注册结果"),
        ("个人资料维护", "已登录用户", "个人资料维护功能允许用户查看和修改自己的基本信息，并支持修改密码。后端通过JWT识别当前用户，不允许前端指定任意用户编号，避免越权修改他人资料。", "ProfileView.vue", "/auth/profile、/auth/password", "AuthService", ["sys_user"], "资料字段、旧密码、新密码", "更新结果"),
    ]),
    ("普通员工模块", [
        ("文本检测", "普通员工", "文本检测是员工提交数据安全检查的主要入口。员工输入任务名称和文本内容，系统识别手机号、身份证号、邮箱、银行卡号、地址和敏感词，生成风险等级、风险评分、命中摘要、脱敏内容和AI建议。", "DetectionView.vue", "/detect/text", "DetectionService、AiReviewService", ["detect_task", "detect_result", "audit_task", "warning_record", "sensitive_word"], "任务名称、文本内容", "检测结果、审核任务"),
        ("文件检测", "普通员工", "文件检测用于处理txt、pdf、doc、docx等文档。后端先抽取文件文本，再进入与文本检测一致的风险识别流程。该功能解决实际业务中员工上传制度文档、合同资料、客户材料前需要检测的问题。", "DetectionView.vue", "/detect/file", "FileTextExtractor、DetectionService", ["detect_task", "detect_result", "audit_task", "warning_record"], "任务名称、文件", "文件检测结果"),
        ("检测结果查看", "普通员工", "检测结果查看功能把系统识别出的风险清晰反馈给员工。页面展示风险等级、风险评分、命中摘要、脱敏结果和AI建议，使员工知道内容为什么存在风险，以及应如何修改。", "DetectionView.vue", "/detect/text、/detect/file", "DetectionService", ["detect_result"], "检测响应", "风险解释信息"),
        ("历史记录查看", "普通员工", "历史记录查看功能用于查询当前员工提交过的检测任务。系统按当前登录用户过滤数据，并关联检测结果和审核任务，使员工能够看到主管审核状态和审核意见。", "HistoryView.vue", "/detect/history", "DetectionService", ["detect_task", "detect_result", "audit_task"], "当前用户", "检测历史列表"),
    ]),
    ("部门主管模块", [
        ("审核申请", "部门主管", "审核申请功能用于处理员工提交后的检测任务。主管查看任务名称、内容类型、风险等级、风险评分、命中摘要、脱敏内容和AI建议，再根据业务要求选择通过或驳回并填写审核意见。", "AuditView.vue", "/audit/task-details、/audit/handle", "AuditService", ["audit_task", "detect_task", "detect_result", "audit_record"], "审核任务、审核意见", "通过或驳回结果"),
        ("部门统计", "部门主管", "部门统计功能用于帮助主管了解部门检测情况。页面展示用户数量、检测次数、风险事件、审核通过率和风险分布等信息，主管可以根据统计结果判断部门数据安全风险趋势。", "ManagerStatsView.vue", "/dashboard/stats、/dashboard/risk-distribution", "DashboardService", ["sys_user", "detect_task", "detect_result", "audit_record"], "统计请求", "统计指标、风险分布"),
        ("风险处理", "部门主管", "风险处理功能用于处理高风险检测任务生成的预警。主管可以查看预警等级、预警内容、关联任务和处理状态，对未处理风险执行处理操作，使风险从发现进入处置阶段。", "WarningView.vue", "/warnings、/warnings/{id}/handle", "WarningService", ["warning_record", "operation_log"], "预警编号", "处理状态"),
        ("审核记录", "部门主管", "审核记录功能用于保存主管的历史审核行为。每次通过或驳回都会形成记录，包括审核任务、审核人、动作、意见和时间。该功能使主管审核具备可追溯性。", "AuditRecordsView.vue", "/audit/records", "AuditService", ["audit_record"], "当前主管身份", "审核记录列表"),
    ]),
    ("数据安全员模块", [
        ("AI文本审核", "数据安全员", "AI文本审核用于对待发布文本进行安全复核。安全员填写审核场景、资料类型、发布范围和文本内容，系统调用DeepSeek或本地兜底策略，返回风险等级、发布建议、风险依据和修改建议，并保存审核记录。", "AiReviewView.vue", "/ai-review", "AiReviewService", ["ai_review_record"], "审核上下文、文本", "AI审核结论"),
        ("AI文件审核", "数据安全员", "AI文件审核用于对合同、财务材料、技术文档等文件进行复核。系统抽取文件文本后调用AI审核服务，并把文件名、内容类型和审核结论保存到AI审核记录表。", "AiReviewView.vue", "/ai-review/file", "FileTextExtractor、AiReviewService", ["ai_review_record"], "审核上下文、文件", "文件审核结论"),
        ("AI审核记录管理", "数据安全员", "AI审核记录管理用于查看历史AI审核。安全员可以查看审核场景、资料类型、发布范围、内容类型、文件名、原文、审核结论和时间，解决审核结果只临时显示的问题。", "AiReviewRecordsView.vue", "/ai-review/records", "AiReviewService", ["ai_review_record"], "当前安全员", "审核记录详情"),
        ("敏感词库管理", "数据安全员", "敏感词库管理用于维护企业自定义风险词。安全员可以新增、查询、修改和删除敏感词，配置分类、风险等级和启用状态。员工检测时会读取启用的敏感词参与风险判断。", "SensitiveWordView.vue", "/sensitive-words", "SensitiveWordService", ["sensitive_category", "sensitive_word"], "敏感词配置", "词库列表"),
        ("风险预警管理", "数据安全员", "风险预警管理用于查看和处理系统生成的高风险预警。安全员根据预警等级和内容判断是否需要进一步处置，并更新预警状态。", "WarningView.vue", "/warnings、/warnings/{id}/handle", "WarningService", ["warning_record", "operation_log"], "预警编号", "预警处理结果"),
        ("风险等级配置", "数据安全员", "风险等级配置用于维护低风险、中风险、高风险和严重风险的分值区间及处理规则。该功能让风险判定规则可以根据管理制度调整。", "RiskLevelView.vue", "/admin/risk-levels", "AdminService", ["risk_level_config"], "等级配置", "保存结果"),
    ]),
    ("系统管理员模块", [
        ("用户管理", "系统管理员", "用户管理用于查看系统用户、按关键字检索用户，并启用或禁用账号。账号禁用后不能继续登录，是管理员控制系统访问的重要手段。", "UserManageView.vue", "/users、/users/{id}/status", "UserManageService", ["sys_user", "operation_log"], "关键字、用户状态", "用户列表"),
        ("角色管理", "系统管理员", "角色管理用于维护普通员工、部门主管、数据安全员和系统管理员等角色信息。角色编码是权限判断的重要依据，角色名称和说明用于后台展示。", "AdminConfigView.vue", "/admin/roles", "AdminService", ["sys_role"], "角色信息", "角色列表"),
        ("权限管理", "系统管理员", "权限管理用于维护菜单或接口权限说明。权限编码和菜单路径说明系统有哪些操作能力，为后续扩展更细粒度RBAC提供基础。", "AdminConfigView.vue", "/admin/permissions", "AdminService", ["sys_permission"], "权限信息", "权限列表"),
        ("日志管理", "系统管理员", "日志管理用于查看登录、检测、审核、预警处理、敏感词维护和配置保存等操作记录。管理员可以根据日志追踪异常操作和业务问题。", "AdminLogView.vue", "/admin/logs", "AdminService、OperationLogService", ["operation_log"], "日志请求", "日志列表"),
        ("系统配置", "系统管理员", "系统配置用于维护JWT有效期、DeepSeek模型名称、文件上传大小等运行参数。该功能减少硬编码配置，提高系统维护灵活性。", "AdminConfigView.vue", "/admin/configs", "AdminService", ["system_config"], "配置项", "配置列表"),
    ]),
]


def feature_records():
    records = []
    for module, items in FEATURES:
        for item in items:
            title, actor, desc, view, api, service, tables, inputs, outputs = item
            records.append({
                "module": module,
                "title": title,
                "actor": actor,
                "desc": desc,
                "view": view,
                "api": api,
                "service": service,
                "tables": tables,
                "inputs": inputs,
                "outputs": outputs,
            })
    return records


def safe_name(text: str) -> str:
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", text)


def draw_center(draw, box, text, fnt=None):
    fnt = fnt or base.font(21)
    base.centered(draw, box, text, fnt, "#000000")


def black_arrow(draw, p1, p2, width=3):
    base.arrow(draw, p1, p2, "#000000", width)


def elbow_arrow(draw, points, width=3):
    """Draw a black orthogonal connector with one arrowhead at the last point."""
    draw.line(points, fill="#000000", width=width, joint="curve")
    if len(points) < 2:
        return
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    angle = math.atan2(y2 - y1, x2 - x1)
    for delta in (2.55, -2.55):
        x = x2 + 16 * math.cos(angle + delta)
        y = y2 + 16 * math.sin(angle + delta)
        draw.line([(x2, y2), (x, y)], fill="#000000", width=width)


def overview_box(draw, box, text, bold=False, size=21):
    draw.rectangle(box, fill="#ffffff", outline="#000000", width=3)
    draw_center(draw, box, text, base.font(size, bold))


def save_black_module_structure():
    path = DIAGRAM_DIR / "00_module_structure_black.png"
    img = Image.new("RGB", (1900, 1050), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "模块结构模型图", font=base.font(34, True), fill="#000000")

    top_box = (760, 125, 1140, 195)
    overview_box(d, top_box, "智能数据守护者系统", True, 23)

    module_names = [module for module, _ in FEATURES]
    module_features = [[item[0] for item in items] for _, items in FEATURES]
    module_w, module_h = 280, 72
    module_y = 315
    xs = [65, 420, 775, 1130, 1485]
    centers = [x + module_w // 2 for x in xs]

    bus_y = 260
    d.line((950, top_box[3], 950, bus_y), fill="#000000", width=3)
    d.line((centers[0], bus_y, centers[-1], bus_y), fill="#000000", width=3)

    for x, cx, module, children in zip(xs, centers, module_names, module_features):
        overview_box(d, (x, module_y, x + module_w, module_y + module_h), module, True, 20)
        elbow_arrow(d, [(cx, bus_y), (cx, module_y)], 3)
        child_y = 430
        prev_bottom = (cx, module_y + module_h)
        for child in children:
            child_box = (x + 15, child_y, x + module_w - 15, child_y + 58)
            overview_box(d, child_box, child, False, 18)
            elbow_arrow(d, [prev_bottom, (cx, child_y)], 2)
            prev_bottom = (cx, child_y + 58)
            child_y += 78

    img.save(path)
    return path


def save_black_logic_relation():
    path = DIAGRAM_DIR / "00_module_logic_black.png"
    img = Image.new("RGB", (1900, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "系统模块主要逻辑关系图", font=base.font(34, True), fill="#000000")

    w, h = 285, 74
    xs = [60, 415, 770, 1125, 1480]
    top_y = 150
    bottom_y = 430
    top_nodes = [
        "用户登录/注册",
        "角色识别与菜单加载",
        "员工提交文本/文件检测",
        "规则识别与AI建议",
        "生成检测结果",
    ]
    bottom_nodes = [
        "管理员维护用户/权限/日志",
        "安全员处理预警/维护规则",
        "高风险生成预警",
        "员工查看审核状态",
        "主管审核通过/驳回",
    ]

    top_boxes = []
    bottom_boxes = []
    for x, text in zip(xs, top_nodes):
        box = (x, top_y, x + w, top_y + h)
        top_boxes.append(box)
        overview_box(d, box, text, False, 20)
    for x, text in zip(xs, bottom_nodes):
        box = (x, bottom_y, x + w, bottom_y + h)
        bottom_boxes.append(box)
        overview_box(d, box, text, False, 20)

    for left, right in zip(top_boxes, top_boxes[1:]):
        y = top_y + h // 2
        elbow_arrow(d, [(left[2], y), (right[0], y)], 3)

    top_last = top_boxes[-1]
    bottom_last = bottom_boxes[-1]
    elbow_arrow(
        d,
        [
            ((top_last[0] + top_last[2]) // 2, top_last[3]),
            ((top_last[0] + top_last[2]) // 2, 335),
            ((bottom_last[0] + bottom_last[2]) // 2, 335),
            ((bottom_last[0] + bottom_last[2]) // 2, bottom_last[1]),
        ],
        3,
    )

    for right, left in zip(reversed(bottom_boxes[1:]), reversed(bottom_boxes[:-1])):
        y = bottom_y + h // 2
        elbow_arrow(d, [(right[0], y), (left[2], y)], 3)

    admin = bottom_boxes[0]
    login = top_boxes[0]
    support_x = login[0] + w // 2
    elbow_arrow(
        d,
        [
            (admin[0] + w // 2, admin[1]),
            (admin[0] + w // 2, 330),
            (support_x, 330),
            (support_x, login[3]),
        ],
        2,
    )
    d.text((90, 335), "基础数据支撑", font=base.font(18), fill="#000000")

    security = bottom_boxes[1]
    rules = top_boxes[3]
    elbow_arrow(
        d,
        [
            (security[0] + w // 2, security[1]),
            (security[0] + w // 2, 285),
            (rules[0] + w // 2, 285),
            (rules[0] + w // 2, rules[3]),
        ],
        2,
    )
    d.text((625, 292), "规则配置反馈", font=base.font(18), fill="#000000")

    d.text((75, 665), "说明：黑色矩形表示业务模块或处理节点，箭头表示数据与任务流向。管理员维护基础权限，安全员维护风险规则，员工、主管和安全员共同形成检测、审核、预警、处置的闭环。", font=base.font(20), fill="#000000")

    img.save(path)
    return path


def draw_process(draw, box, text):
    draw.rectangle(box, fill="#ffffff", outline="#000000", width=3)
    draw_center(draw, box, text, base.font(20))


def draw_terminal(draw, box, text):
    draw.rounded_rectangle(box, radius=35, fill="#ffffff", outline="#000000", width=3)
    draw_center(draw, box, text, base.font(20, True))


def draw_io(draw, box, text):
    x1, y1, x2, y2 = box
    slant = 34
    points = [(x1 + slant, y1), (x2, y1), (x2 - slant, y2), (x1, y2)]
    draw.polygon(points, fill="#ffffff", outline="#000000")
    draw.line(points + [points[0]], fill="#000000", width=3)
    draw_center(draw, box, text, base.font(20))


def draw_decision(draw, cx, cy, w, h, text):
    points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(points, fill="#ffffff", outline="#000000")
    draw.line(points + [points[0]], fill="#000000", width=3)
    draw_center(draw, (cx - w // 2 + 12, cy - h // 2 + 12, cx + w // 2 - 12, cy + h // 2 - 12), text, base.font(18))


def steps_for(feature):
    title = feature["title"]
    mapping = {
        "验证码获取": ["打开登录页", "请求验证码", "生成验证码", "返回图片", "用户输入验证码"],
        "用户登录": ["填写登录信息", "提交登录", "校验验证码", "校验账号密码", "生成JWT", "跳转角色首页"],
        "用户注册": ["填写注册资料", "提交注册", "检查用户名", "加密密码", "保存用户", "返回结果"],
        "个人资料维护": ["进入个人中心", "加载资料", "编辑资料或密码", "提交保存", "校验身份", "更新用户"],
        "文本检测": ["填写任务和文本", "提交检测", "规则识别", "敏感词匹配", "计算风险", "生成结果和审核任务"],
        "文件检测": ["选择检测文件", "上传文件", "抽取文本", "执行检测", "生成结果", "进入审核"],
        "检测结果查看": ["接收检测结果", "展示风险等级", "展示命中摘要", "展示脱敏内容", "展示AI建议"],
        "历史记录查看": ["进入历史页", "查询个人任务", "关联结果和审核", "展示历史列表", "查看审核意见"],
        "审核申请": ["进入审核页", "加载待审任务", "查看风险详情", "填写审核意见", "通过或驳回", "写入审核记录"],
        "部门统计": ["进入统计页", "请求统计数据", "计算检测数量", "计算风险分布", "展示统计图表"],
        "风险处理": ["进入风险页面", "加载预警列表", "查看预警内容", "选择未处理项", "更新处理状态"],
        "审核记录": ["进入记录页", "查询审核记录", "按时间排序", "展示动作和意见"],
        "AI文本审核": ["填写审核上下文", "输入文本", "提交AI审核", "调用模型或兜底", "保存审核记录", "展示结论"],
        "AI文件审核": ["上传审核文件", "抽取文件文本", "提交AI审核", "保存文件记录", "展示结论"],
        "AI审核记录管理": ["进入记录页", "查询AI审核记录", "展示记录列表", "打开详情", "查看原文和结论"],
        "敏感词库管理": ["进入词库页", "查询词条", "新增或编辑", "校验分类", "保存词条"],
        "风险预警管理": ["进入预警页", "查询预警", "查看风险内容", "处理预警", "更新状态"],
        "风险等级配置": ["进入配置页", "加载等级", "编辑分值规则", "校验区间", "保存配置"],
        "用户管理": ["进入用户页", "查询用户", "筛选用户", "启用或禁用", "刷新列表"],
        "角色管理": ["进入角色页", "查询角色", "编辑角色", "保存角色", "刷新列表"],
        "权限管理": ["进入权限页", "查询权限", "编辑权限", "保存权限", "刷新列表"],
        "日志管理": ["进入日志页", "查询日志", "按时间展示", "查看操作详情"],
        "系统配置": ["进入配置页", "查询配置", "编辑配置值", "保存配置", "刷新列表"],
    }
    return mapping.get(title, ["进入页面", "填写数据", "提交请求", "后端处理", "返回结果"])


def decision_for(feature):
    title = feature["title"]
    if "登录" in title:
        return "校验通过？", "进入系统", "提示失败"
    if "注册" in title:
        return "用户名可用？", "保存用户", "提示重复"
    if "文件" in title:
        return "文件可解析？", "继续处理", "提示重传"
    if "审核" in title and title != "AI审核记录管理":
        return "审核通过？", "通过/保存", "驳回/提示"
    if "配置" in title or "管理" in title:
        return "数据有效？", "保存/展示", "提示修改"
    if "风险" in title:
        return "需处理？", "更新状态", "保持记录"
    return "处理成功？", "返回结果", "提示错误"


def flow_chart(feature, idx):
    path = DIAGRAM_DIR / f"{idx:02d}_{safe_name(feature['title'])}_flow_chart.png"
    img = Image.new("RGB", (1200, 1250), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((55, 52), f"流程图：{feature['title']}", font=base.font(30, True), fill="#000000")
    d.text((70, 122), f"参与者：{feature['actor']}", font=base.font(21), fill="#000000")
    draw_terminal(d, (460, 155, 740, 220), "开始")
    prev = (600, 220)
    y = 260
    steps = steps_for(feature)
    decision_index = max(1, min(3, len(steps) - 2))
    q, yes, no = decision_for(feature)
    for i, step in enumerate(steps):
        x1 = 420
        if i == 0 or "填写" in step or "上传" in step or "选择" in step or "输入" in step:
            draw_io(d, (x1, y, x1 + 360, y + 66), step)
        else:
            draw_process(d, (x1, y, x1 + 360, y + 66), step)
        black_arrow(d, prev, (600, y))
        prev = (600, y + 66)
        y += 95
        if i == decision_index:
            draw_decision(d, 600, y + 58, 320, 116, q)
            black_arrow(d, prev, (600, y))
            yes_box = (800, y + 28, 1085, y + 88)
            no_box = (115, y + 28, 400, y + 88)
            draw_process(d, yes_box, yes)
            draw_process(d, no_box, no)
            black_arrow(d, (760, y + 58), (800, y + 58))
            black_arrow(d, (440, y + 58), (400, y + 58))
            d.text((770, y + 30), "是", font=base.font(16, True), fill="#000000")
            d.text((415, y + 30), "否", font=base.font(16, True), fill="#000000")
            yes_end = (820, y + 155, 1065, y + 215)
            no_end = (135, y + 155, 380, y + 215)
            draw_terminal(d, yes_end, "结束")
            draw_terminal(d, no_end, "结束")
            black_arrow(d, ((yes_box[0] + yes_box[2]) // 2, yes_box[3]), ((yes_end[0] + yes_end[2]) // 2, yes_end[1]), 2)
            black_arrow(d, ((no_box[0] + no_box[2]) // 2, no_box[3]), ((no_end[0] + no_end[2]) // 2, no_end[1]), 2)
            img.save(path)
            return path
    draw_terminal(d, (460, y + 10, 740, y + 75), "结束")
    black_arrow(d, prev, (600, y + 10))
    img.save(path)
    return path


def business_model(feature, idx):
    path = DIAGRAM_DIR / f"{idx:02d}_{safe_name(feature['title'])}_business_model.png"
    img = Image.new("RGB", (1500, 780), "#ffffff")
    d = ImageDraw.Draw(img)
    d.text((55, 52), f"业务模型图：{feature['title']}", font=base.font(30, True), fill="#000000")
    boxes = [
        ("角色", feature["actor"], 70),
        ("页面", feature["view"], 320),
        ("接口", feature["api"], 570),
        ("服务", feature["service"], 820),
        ("数据表", "、".join(feature["tables"]), 1070),
    ]
    y = 180
    for label, value, x in boxes:
        d.rectangle((x, y, x + 200, y + 210), outline="#000000", width=3)
        d.line((x, y + 44, x + 200, y + 44), fill="#000000", width=2)
        draw_center(d, (x, y, x + 200, y + 44), label, base.font(19, True))
        draw_center(d, (x + 8, y + 56, x + 192, y + 198), value, base.font(17))
    for i in range(4):
        black_arrow(d, (270 + i * 250, y + 105), (320 + i * 250, y + 105), 3)
    d.rectangle((110, 500, 640, 610), outline="#000000", width=3)
    d.line((110, 545, 640, 545), fill="#000000", width=2)
    draw_center(d, (110, 500, 640, 545), "输入", base.font(20, True))
    draw_center(d, (120, 555, 630, 600), feature["inputs"], base.font(18))
    d.rectangle((820, 500, 1370, 610), outline="#000000", width=3)
    d.line((820, 545, 1370, 545), fill="#000000", width=2)
    draw_center(d, (820, 500, 1370, 545), "输出", base.font(20, True))
    draw_center(d, (830, 555, 1360, 600), feature["outputs"], base.font(18))
    black_arrow(d, (640, 555), (820, 555), 2)
    d.text((110, 665), "说明：角色通过页面提交输入，接口接收请求，服务完成校验和业务处理，数据表保存或提供业务数据，最后返回输出结果。", font=base.font(19), fill="#000000")
    img.save(path)
    return path


def generate_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    result = {}
    for idx, feature in enumerate(feature_records(), 1):
        result[feature["title"]] = (flow_chart(feature, idx), business_model(feature, idx))
    return result


def plain_detail(feature):
    title = feature["title"]
    paragraphs = [
        f"{title}由{feature['actor']}使用，页面入口为{feature['view']}，后端接口为{feature['api']}。该功能的输入是{feature['inputs']}，输出是{feature['outputs']}。设计时前端负责收集和展示数据，后端负责权限校验、参数校验、业务处理和结果返回。",
        f"该功能的后端处理主要由{feature['service']}完成，涉及数据表包括{'、'.join(feature['tables'])}。如果功能需要保存业务记录，服务层在返回结果前完成落库；如果功能是查询类操作，服务层按当前登录用户或角色权限过滤数据，避免越权访问。",
        f"流程图使用开始/结束框、处理框、输入输出框、判断菱形和箭头表示执行顺序。判断节点保留成功和失败两条路径，使正常处理和异常提示都能在图中体现。",
        f"业务模型图说明该功能从角色到页面、接口、服务和数据表的调用关系。该模型用于解释功能不是单纯的页面展示，而是完整的前后端业务链路。测试时需要同时验证正常提交、错误输入、权限限制和数据保存结果。",
        f"{title}的业务规则要求前端字段名称、接口参数名称和后端实体字段保持一致。用户提交数据后，系统先进行基本校验，再进入业务处理。处理成功时返回统一响应结构，处理失败时返回明确错误信息，前端根据响应结果给出提示。",
        f"{title}还需要考虑和其他模块的衔接。检测类功能会影响审核任务和预警记录，审核类功能会影响员工历史记录，安全治理类功能会影响后续检测规则，管理员配置类功能会影响系统运行参数。因此详细设计不能只描述按钮点击，还要说明数据如何继续流转。",
        f"实现该功能时，前端需要保持加载状态、错误提示和空数据状态完整，避免用户重复提交或误以为系统无响应。后端需要把参数校验放在业务处理之前，并保持事务边界清晰，涉及多张表写入时要么全部成功，要么全部回滚。",
        f"该功能的边界条件包括无权限访问、输入为空、关联数据不存在、网络请求失败和数据库保存失败。设计文档将这些情况放入流程图的判断路径中，目的在于说明系统不仅能处理正常流程，也能处理失败和异常流程。",
    ]
    specific = {
        "验证码获取": ["验证码生成后应具有有效期，登录校验成功或过期后应失效。验证码不能只在前端判断，必须由后端统一校验。", "测试时需要刷新验证码、输入错误验证码、输入过期验证码和输入正确验证码，确认登录流程能正确区分不同情况。"],
        "用户登录": ["登录成功后前端保存令牌和用户角色，后续请求通过请求拦截器携带Authorization头。后端从令牌中解析用户身份，供业务服务获取当前用户。", "登录失败时不应暴露过多细节，例如不应明确提示用户名存在但密码错误，以免被用于账号枚举。"],
        "用户注册": ["注册用户默认应按普通权限处理，不能让用户自行注册为系统管理员或数据安全员。密码保存前必须加密，接口返回结果中不能包含原始密码。", "注册功能在课程设计中用于说明系统扩展能力，演示账号则由初始化程序创建，二者互不冲突。"],
        "个人资料维护": ["修改资料时后端必须使用当前登录用户编号，不能信任前端传入任意用户ID。修改密码时必须校验旧密码，旧密码错误时直接拒绝更新。", "该功能虽然不是核心检测流程，但它保证用户资料准确，为后续日志审计和责任追踪提供基础。"],
        "文本检测": ["文本检测完成后会同时写入检测任务和检测结果，并自动创建主管审核任务。若风险等级较高，还会生成预警记录。", "该功能测试时应准备包含手机号、邮箱、身份证号、银行卡号和敏感词的文本，确认命中摘要和脱敏结果正确。"],
        "文件检测": ["文件检测先抽取文本，再复用文本检测规则，保证文本输入和文件输入使用同一套风险判断标准。文件解析失败时不应保存无效检测任务。", "该功能测试时需要覆盖txt、pdf、doc、docx等类型，至少验证可解析文件和不可解析文件两类情况。"],
        "检测结果查看": ["检测结果应突出风险等级和风险原因，员工能够根据命中摘要理解哪些字段需要修改。脱敏内容用于降低二次泄露风险。", "结果页面只展示当前提交任务的检测结论，历史结果则进入历史记录页面统一查询。"],
        "历史记录查看": ["历史记录必须按当前员工过滤，员工不能查看其他人的检测任务。记录中同时展示检测状态和审核状态，便于员工确认流程进度。", "主管审核通过或驳回后，员工刷新历史记录即可看到审核意见，这体现了员工端和主管端的数据联动。"],
        "审核申请": ["主管审核时需要同时参考原始内容、脱敏内容、风险等级、风险评分、命中摘要和AI建议。通过或驳回都必须填写或保存审核意见。", "审核动作会更新审核任务状态，也会同步更新检测任务状态，员工历史记录由此显示最终审核结果。"],
        "部门统计": ["部门统计不直接改变业务数据，而是读取用户、检测、结果和审核记录形成统计指标。统计结果帮助主管判断部门整体风险情况。", "测试时需要在系统中先产生若干检测和审核记录，再查看统计页面是否能正确展示数量和分布。"],
        "风险处理": ["风险处理面向高风险或严重风险预警。主管或安全员处理预警后，系统更新预警状态，并通过日志记录处理动作。", "该功能不能简单删除预警，因为预警记录本身属于安全审计数据，应保留处理状态和处理时间。"],
        "审核记录": ["审核记录是主管工作的追溯依据。每次审核都应记录审核任务编号、审核人、动作、意见和时间。", "审核记录页面主要用于查询，不应允许随意修改历史审核动作，否则会破坏审核可信度。"],
        "AI文本审核": ["AI文本审核的提示词应包含审核场景、资料类型和发布范围，这些上下文会影响模型判断。审核结论保存到记录表，便于复盘。", "如果DeepSeek调用失败，系统返回本地兜底建议，保证页面仍有结果，避免演示流程中断。"],
        "AI文件审核": ["AI文件审核与员工文件检测不同，它侧重发布决策建议，而不是规则评分。文件名和文件内容类型需要随审核结论一起保存。", "测试时应上传包含客户资料、报价、手机号或接口密钥的文件，确认审核结论能指出风险依据。"],
        "AI审核记录管理": ["AI审核记录管理让安全员能查看过去的文本审核和文件审核。详情中展示原文和结论，便于解释某次判断依据。", "该功能提升了系统可追溯性，也证明AI审核不是一次性临时输出，而是系统安全治理数据的一部分。"],
        "敏感词库管理": ["敏感词库保存企业自定义风险词，弥补正则规则无法识别业务语义的不足。词条启用后会参与后续检测。", "保存敏感词前要校验分类存在，风险等级要与系统支持的等级一致，避免检测评分出现异常。"],
        "风险预警管理": ["风险预警由高风险检测自动生成，安全员查看后进行处理。处理动作改变预警状态，但不删除预警记录。", "该功能和风险处理功能使用相同预警数据，但数据安全员更关注安全治理闭环，主管更关注部门业务处置。"],
        "风险等级配置": ["风险等级配置定义分值区间和处理规则。低、中、高、严重风险应覆盖完整分值范围，不能交叉或遗漏。", "保存配置后，页面应刷新展示最新规则，后续检测说明也可以引用这些处理规则。"],
        "用户管理": ["管理员通过用户管理控制账号状态。禁用账号后，该用户不能继续登录，从运维层面阻断访问。", "该功能需要写入操作日志，因为启用或禁用账号属于影响系统访问的重要动作。"],
        "角色管理": ["角色编码是权限判断的关键标识，普通员工、部门主管、数据安全员和系统管理员分别对应不同菜单和接口权限。", "角色管理主要维护角色名称和说明，角色编码不宜随意修改，否则会影响路由守卫和后端权限判断。"],
        "权限管理": ["权限管理保存权限编码、权限名称和菜单路径，用于说明系统中有哪些菜单或接口能力。", "当前系统以角色控制为主，权限表为后续扩展角色权限绑定提供数据基础。"],
        "日志管理": ["日志管理集中展示系统关键操作，包括登录、检测、审核、敏感词维护、预警处理和配置保存。", "管理员可以根据日志中的用户、类型、时间和详情定位异常操作，是系统审计的重要依据。"],
        "系统配置": ["系统配置用于维护可变参数，例如JWT有效期、DeepSeek模型名称和文件上传限制。", "配置值保存前应进行基本校验，关键配置修改后应通过日志记录，便于出现问题时回溯。"],
    }
    return paragraphs + specific.get(title, [])


def add_function(doc, feature, diagrams, fig_no):
    doc.add_heading(feature["title"], level=3)
    base.add_para(doc, feature["desc"])
    for p in plain_detail(feature):
        base.add_para(doc, p)
    base.add_table(doc, ["设计项", "内容"], [
        ["参与者", feature["actor"]],
        ["前端页面", feature["view"]],
        ["后端接口", feature["api"]],
        ["业务服务", feature["service"]],
        ["数据表", "、".join(feature["tables"])],
        ["输入", feature["inputs"]],
        ["输出", feature["outputs"]],
    ], widths=[3.0, 12.5])
    activity, model = diagrams[feature["title"]]
    base.add_picture(doc, activity, f"图2-{fig_no[0]} {feature['title']}流程图", 5.7)
    fig_no[0] += 1
    base.add_picture(doc, model, f"图2-{fig_no[0]} {feature['title']}业务模型图", 5.9)
    fig_no[0] += 1


def add_toc(doc):
    doc.add_heading("目录", level=1)
    lines = [
        "1 可行性研究与需求分析\t1",
        "1.1 选题意义\t1",
        "1.2 可行性研究\t2",
        "1.3 需求分析\t3",
        "2 系统设计\t4",
        "2.1 设计原则\t4",
        "2.2 主要模块功能（概要设计）\t5",
        "2.3 各个模块详细设计\t7",
        "2.3.1 登录注册模块\t7",
        "2.3.2 普通员工模块\t12",
        "2.3.3 部门主管模块\t17",
        "2.3.4 数据安全员模块\t22",
        "2.3.5 系统管理员模块\t30",
        "2.4 数据库概念结构模型\t36",
        "3 系统实现与测试\t40",
        "4 总结\t48",
    ]
    for line in lines:
        doc.add_paragraph(line)
    doc.add_page_break()


def add_common_sections(doc, overview):
    doc.add_heading("1 可行性研究与需求分析", level=1)
    doc.add_heading("1.1 选题意义", level=2)
    for text in [
        "智能数据守护者系统面向企业内容发布和内部资料流转中的数据安全问题。员工在日常工作中会处理客户名单、合同报价、财务报表、技术文档、接口密钥和个人隐私信息，如果缺少发布前检测与审核机制，容易造成隐私泄露和商业秘密泄露。",
        "本系统把敏感信息检测、主管审核、AI辅助审核、风险预警、敏感词维护、权限管理和日志审计放在统一平台中。普通员工负责提交文本或文件，部门主管负责审核员工检测任务，数据安全员负责规则治理和AI复核，系统管理员负责用户、角色、权限和配置维护。",
        "课题覆盖Spring Boot、MyBatis Plus、MySQL、Vue3、Element Plus、JWT、文件上传、DeepSeek接口、流程图、业务模型和E-R建模等内容，能够体现课程设计对需求分析、系统设计、数据库设计、编码实现和测试展示的综合要求。",
    ]:
        base.add_para(doc, text)
    doc.add_heading("1.2 可行性研究", level=2)
    for heading, text in [
        ("1.2.1 技术可行性", "后端采用Spring Boot和MyBatis Plus，前端采用Vue3、Vite和Element Plus，数据库采用MySQL。技术栈成熟，资料丰富，适合在课程设计周期内完成。AI审核通过DeepSeek接口实现，同时保留本地兜底建议，避免外部接口异常导致流程中断。"),
        ("1.2.2 运行可行性", "系统可在普通开发机中运行，后端通过Maven启动，前端通过Vite启动，数据库使用本地MySQL。系统主要处理文本、文件解析和数据库读写，对硬件要求较低，适合课堂演示和本地测试。"),
        ("1.2.3 操作可行性", "系统按角色划分菜单，员工、主管、安全员和管理员进入系统后看到不同功能。页面使用表单、表格、上传控件和弹窗组织操作，用户按照业务流程即可完成检测、审核、治理和管理。"),
        ("1.2.4 法律可行性", "系统目标是降低敏感数据泄露风险，符合访问控制、最小必要、操作留痕和人工复核原则。课程演示不使用真实企业数据，密码加密存储，敏感内容提供脱敏展示，AI审核只作为辅助建议。"),
    ]:
        doc.add_heading(heading, level=3)
        base.add_para(doc, text)
    doc.add_heading("1.3 需求分析", level=2)
    for heading, items in [
        ("1.3.1 安全性需求", ["登录必须校验验证码、账号状态和密码。", "接口必须按JWT和角色进行访问控制。", "检测结果需要提供脱敏展示。", "关键业务操作需要写入日志。", "高风险任务需要触发预警并支持处理。"]),
        ("1.3.2 稳定性和可维护性需求", ["系统采用Controller、Service、Mapper、Entity分层结构。", "AI接口不可用时返回本地兜底建议。", "敏感词、风险等级和系统配置支持后台维护。", "数据库记录保留创建时间、更新时间和逻辑删除标记。"]),
        ("1.3.3 主要功能需求", ["登录注册模块：验证码、登录、注册、个人资料维护。", "普通员工模块：文本检测、文件检测、检测结果查看、历史记录查看。", "部门主管模块：审核申请、部门统计、风险处理、审核记录。", "数据安全员模块：AI文本审核、AI文件审核、AI审核记录、敏感词库、风险预警、风险等级配置。", "系统管理员模块：用户管理、角色管理、权限管理、日志管理、系统配置。"]),
    ]:
        doc.add_heading(heading, level=3)
        base.add_bullets(doc, items)
    doc.add_heading("2 系统设计", level=1)
    doc.add_heading("2.1 设计原则", level=2)
    base.add_para(doc, "系统设计遵循角色分工清晰、业务流程闭环、数据可追溯、规则可配置和异常可处理原则。员工提交检测，主管审核，安全员维护安全规则，管理员维护系统基础数据。检测结果、审核记录、AI审核记录、预警记录和操作日志共同保证业务过程可追踪。")
    base.add_picture(doc, overview["system_data_model"], "图2-1 系统数据模型图", 6.2)
    doc.add_heading("2.2 主要模块功能（概要设计）", level=2)
    base.add_para(doc, "系统由登录注册、普通员工、部门主管、数据安全员、系统管理员和公共支撑模块组成。公共支撑模块包括JWT认证、统一响应、异常处理、跨域配置、文件文本抽取和操作日志。各角色模块通过REST接口连接前端页面和后端服务，后端服务再访问数据库。")
    base.add_picture(doc, overview["module_structure"], "图2-2 模块结构模型图", 6.2)
    base.add_picture(doc, overview["module_logic"], "图2-3 系统模块主要逻辑关系图", 6.2)


def add_detailed_design(doc, diagrams):
    doc.add_heading("2.3 各个模块详细设计", level=2)
    fig_no = [4]
    section_no = 1
    for module, _ in FEATURES:
        doc.add_heading(f"2.3.{section_no} {module}", level=3)
        intro = {
            "登录注册模块": "登录注册模块是系统访问入口，负责验证码、登录、注册和个人资料维护。",
            "普通员工模块": "普通员工模块是检测任务入口，负责提交文本或文件并查看检测和审核结果。",
            "部门主管模块": "部门主管模块负责审核员工检测任务、查看部门统计、处理风险和追踪审核记录。",
            "数据安全员模块": "数据安全员模块负责AI审核、审核记录、敏感词库、风险预警和风险等级配置。",
            "系统管理员模块": "系统管理员模块负责系统运维，维护用户、角色、权限、日志和配置。",
        }[module]
        base.add_para(doc, intro)
        for feature in [x for x in feature_records() if x["module"] == module]:
            add_function(doc, feature, diagrams, fig_no)
        section_no += 1


def add_database_test_summary(doc, overview):
    doc.add_heading("2.4 数据库概念结构模型", level=2)
    base.add_para(doc, "数据库分为用户权限、检测审核、安全治理和系统运维四类数据。用户权限类表保存用户、角色和权限；检测审核类表保存检测任务、检测结果、审核任务和审核记录；安全治理类表保存AI审核记录、敏感词、风险等级和预警；运维类表保存配置和日志。")
    doc.add_heading("2.4.1 表结构设计（数据库表结构表格）", level=3)
    tables = base.parse_schema()
    important = ["sys_user", "sys_role", "sys_permission", "detect_task", "detect_result", "audit_task", "audit_record", "warning_record", "ai_review_record", "sensitive_word", "risk_level_config", "system_config", "operation_log"]
    idx = 1
    for table in important:
        if table in tables:
            p = doc.add_paragraph(f"表2-{idx} {table}表结构")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            base.add_table(doc, ["字段名", "类型", "是否可空", "字段说明"], [list(r) for r in tables[table]], widths=[3.5, 3.0, 2.0, 7.0])
            idx += 1
    doc.add_heading("2.4.2 E-R模型", level=3)
    base.add_picture(doc, overview["er"], "图2-52 E-R模型图", 6.2)
    doc.add_heading("3 系统实现与测试", level=1)
    base.add_para(doc, "系统实现采用前后端分离架构。前端通过Vue3和Element Plus实现页面，后端通过Spring Boot提供接口，MyBatis Plus访问MySQL。测试使用employee、manager、security、admin四个账号，分别验证员工、主管、安全员和管理员功能。")
    base.add_table(doc, ["测试项", "测试步骤", "预期结果", "结论"], [
        ["员工检测", "提交文本和文件检测", "生成检测结果和审核任务", "通过"],
        ["主管审核", "查看审核申请并通过或驳回", "员工历史记录显示审核结果", "通过"],
        ["部门统计", "打开统计页面", "显示检测数量和风险分布", "通过"],
        ["AI审核", "提交文本和文件审核", "返回结论并保存记录", "通过"],
        ["预警处理", "处理高风险预警", "状态更新为已处理", "通过"],
        ["管理员维护", "查看用户、角色、权限、日志、配置", "页面数据正常展示", "通过"],
    ], widths=[3.0, 5.0, 6.0, 2.0])
    shots = [
        ("01-login.png", "图3-1 登录界面"),
        ("02-employee-detect.png", "图3-2 员工检测界面"),
        ("03-employee-history.png", "图3-3 员工历史记录界面"),
        ("05-manager-audit.png", "图3-4 主管审核申请界面"),
        ("06-manager-stats.png", "图3-5 主管部门统计界面"),
        ("07-manager-risk.png", "图3-6 主管风险处理界面"),
        ("08-manager-records.png", "图3-7 主管审核记录界面"),
        ("09-security-ai.png", "图3-8 安全员AI审核界面"),
        ("10-security-sensitive.png", "图3-9 敏感词库界面"),
        ("13-admin-users.png", "图3-10 管理员用户管理界面"),
        ("15-admin-permissions.png", "图3-11 管理员权限管理界面"),
        ("16-admin-logs.png", "图3-12 管理员日志界面"),
    ]
    for file, cap in shots:
        base.add_picture(doc, SCREENSHOT_DIR / file, cap, 6.2)
    doc.add_heading("4 总结", level=1)
    for text in [
        "本系统完成了智能数据守护者从检测、审核、AI复核、预警处理到系统管理的完整流程。普通员工提交数据，部门主管完成审核，数据安全员维护安全规则并进行AI复核，系统管理员维护基础数据和日志。",
        "本版报告补齐了部门主管模块，并将详细设计按每个功能展开。每个功能都给出说明、设计项、流程图和业务模型图，能够直接说明功能如何从页面、接口、服务落实到数据库。",
        "后续系统可以继续完善分页查询、部门数据范围控制、AI审核结构化存储、更多文件格式解析和自动化测试用例，使系统更接近真实企业数据安全治理平台。",
    ]:
        base.add_para(doc, text)


def count_cn(doc):
    return len(re.findall(r"[\u4e00-\u9fff]", "\n".join(p.text for p in doc.paragraphs)))


def build():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = generate_diagrams()
    requested_dir = REPORT_DIR / "assets" / "requested-template-diagrams"
    overview = {
        "system_data_model": requested_dir / "01_system_data_model.png",
        "module_structure": save_black_module_structure(),
        "module_logic": save_black_logic_relation(),
        "er": requested_dir / "08_er_model.png",
    }
    doc = Document()
    base.configure_doc(doc)
    base.add_cover(doc)
    add_toc(doc)
    add_common_sections(doc, overview)
    add_detailed_design(doc, diagrams)
    add_database_test_summary(doc, overview)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("智能数据守护者系统课程设计报告")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    doc.save(OUT)
    print(OUT)
    print(count_cn(doc))
    print(len(doc.inline_shapes))
    print(len(doc.tables))


if __name__ == "__main__":
    build()
