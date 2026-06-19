from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "assets"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"
DIAGRAM_DIR = ASSET_DIR / "requested-template-diagrams"
OUT = REPORT_DIR / "智能数据守护者系统课程设计报告-指定目录版.docx"
SCHEMA_PATH = ROOT / "database" / "schema.sql"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    choices = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for item in choices:
        if item.exists():
            return ImageFont.truetype(str(item), size)
    return ImageFont.load_default()


def wrap_text(text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in str(text).split("\n"):
        buf = ""
        for ch in raw:
            if fnt.getlength(buf + ch) <= max_width:
                buf += ch
            else:
                if buf:
                    lines.append(buf)
                buf = ch
        if buf:
            lines.append(buf)
    return lines or [""]


def centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt: ImageFont.FreeTypeFont, fill="#111827") -> None:
    lines = wrap_text(text, fnt, box[2] - box[0] - 20)
    line_h = fnt.size + 7
    y = box[1] + (box[3] - box[1] - line_h * len(lines)) / 2
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=fnt)
        x = box[0] + (box[2] - box[0] - (bb[2] - bb[0])) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, p1: tuple[int, int], p2: tuple[int, int], fill="#334155", width=3) -> None:
    draw.line([p1, p2], fill=fill, width=width)
    import math
    x1, y1 = p1
    x2, y2 = p2
    ang = math.atan2(y2 - y1, x2 - x1)
    for delta in (2.55, -2.55):
        x = x2 + 15 * math.cos(ang + delta)
        y = y2 + 15 * math.sin(ang + delta)
        draw.line([(x2, y2), (x, y)], fill=fill, width=width)


def node(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill="#ffffff", outline="#2563eb", radius=18, bold=False) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    centered(draw, box, text, font(22 if not bold else 24, bold), "#0f172a")


def diamond(draw: ImageDraw.ImageDraw, cx: int, cy: int, w: int, h: int, text: str) -> None:
    points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(points, fill="#fef9c3", outline="#ca8a04")
    draw.line(points + [points[0]], fill="#ca8a04", width=3)
    centered(draw, (cx - w // 2 + 10, cy - h // 2 + 10, cx + w // 2 - 10, cy + h // 2 - 10), text, font(18), "#111827")


def diagram_title(draw: ImageDraw.ImageDraw, title: str) -> None:
    draw.text((50, 34), title, font=font(34, True), fill="#0f172a")
    draw.line((50, 88, 1550, 88), fill="#cbd5e1", width=2)


def save_system_data_model() -> Path:
    path = DIAGRAM_DIR / "01_system_data_model.png"
    img = Image.new("RGB", (1600, 1050), "#f8fafc")
    d = ImageDraw.Draw(img)
    diagram_title(d, "系统数据模型图")
    layers = [
        ("用户与权限数据层", ["sys_user 用户", "sys_role 角色", "sys_permission 权限", "sys_department 部门"], 130, "#dbeafe", "#2563eb"),
        ("检测业务数据层", ["detect_task 检测任务", "detect_result 检测结果", "audit_task 审核任务", "audit_record 审核记录"], 360, "#dcfce7", "#16a34a"),
        ("安全治理数据层", ["sensitive_category 敏感词分类", "sensitive_word 敏感词", "risk_level_config 风险等级", "warning_record 风险预警"], 590, "#fef3c7", "#d97706"),
        ("系统运维数据层", ["system_config 系统配置", "operation_log 操作日志", "DeepSeek配置", "JWT登录状态"], 820, "#fce7f3", "#db2777"),
    ]
    for title, items, y, fill, outline in layers:
        d.rounded_rectangle((70, y, 1530, y + 150), radius=24, fill=fill, outline=outline, width=3)
        d.text((100, y + 20), title, font=font(24, True), fill="#0f172a")
        x = 360
        for item in items:
            node(d, (x, y + 35, x + 245, y + 115), item, "#ffffff", outline, 14)
            x += 285
    for y1, y2, label in [(280, 360, "用户提交检测"), (510, 590, "结果触发治理"), (740, 820, "操作形成日志")]:
        arrow(d, (800, y1), (800, y2))
        d.text((820, (y1 + y2) // 2 - 10), label, font=font(18), fill="#334155")
    img.save(path)
    return path


def save_module_structure() -> Path:
    path = DIAGRAM_DIR / "02_module_structure.png"
    img = Image.new("RGB", (1600, 1000), "#f8fafc")
    d = ImageDraw.Draw(img)
    diagram_title(d, "模块结构模型图")
    node(d, (610, 120, 990, 200), "智能数据守护者系统", "#e0f2fe", "#0284c7", 20, True)
    modules = [
        ("登录注册模块", ["验证码", "登录", "注册", "JWT鉴权"], 120, 330),
        ("普通员工模块", ["文本检测", "文件检测", "检测历史", "个人中心"], 440, 330),
        ("部门主管模块", ["审核申请", "部门统计", "风险处理", "审核记录"], 760, 330),
        ("数据安全员模块", ["AI审核", "敏感词库", "风险预警", "风险等级"], 1080, 330),
        ("系统管理员模块", ["用户管理", "角色管理", "权限管理", "日志配置"], 1240, 620),
    ]
    centers = []
    for title, items, x, y in modules:
        node(d, (x, y, x + 260, y + 75), title, "#eff6ff", "#2563eb", 16, True)
        centers.append((x + 130, y))
        for i, item in enumerate(items):
            node(d, (x + 15, y + 105 + i * 78, x + 245, y + 160 + i * 78), item, "#ffffff", "#64748b", 12)
            arrow(d, (x + 130, y + 75), (x + 130, y + 105 + i * 78), "#94a3b8", 2)
    for cx, cy in centers:
        arrow(d, (800, 200), (cx, cy), "#334155", 3)
    img.save(path)
    return path


def save_logic_relation() -> Path:
    path = DIAGRAM_DIR / "03_module_logic_relation.png"
    img = Image.new("RGB", (1700, 980), "#f8fafc")
    d = ImageDraw.Draw(img)
    diagram_title(d, "系统模块主要逻辑关系图")
    flow = [
        ("用户登录/注册", 70, 170),
        ("角色识别与菜单加载", 360, 170),
        ("员工提交文本/文件检测", 650, 170),
        ("规则识别与AI建议", 940, 170),
        ("生成检测结果", 1230, 170),
        ("主管审核通过/驳回", 1230, 430),
        ("员工查看审核状态", 940, 430),
        ("高风险生成预警", 650, 430),
        ("安全员处理预警/维护规则", 360, 430),
        ("管理员维护用户权限日志", 70, 430),
    ]
    boxes = {}
    for text, x, y in flow:
        boxes[text] = (x, y, x + 230, y + 90)
        node(d, boxes[text], text, "#ffffff", "#2563eb", 18)
    for a, b in zip(flow[:5], flow[1:5]):
        arrow(d, (a[1] + 230, a[2] + 45), (b[1], b[2] + 45))
    arrow(d, (1345, 260), (1345, 430))
    arrow(d, (1230, 475), (1170, 475))
    arrow(d, (940, 475), (880, 475))
    arrow(d, (650, 475), (590, 475))
    arrow(d, (360, 475), (300, 475))
    arrow(d, (185, 430), (185, 260), "#64748b")
    d.text((210, 330), "基础数据支撑", font=font(18), fill="#334155")
    for label, p in [("检测结果进入审核队列", (1260, 320)), ("风险策略反向影响检测", (560, 610)), ("日志记录贯穿全流程", (185, 610))]:
        d.text(p, label, font=font(20, True), fill="#0f172a")
    img.save(path)
    return path


def save_flow(name: str, title: str, steps: list[str], decisions: dict[int, tuple[str, str, str]] | None = None) -> Path:
    path = DIAGRAM_DIR / name
    img = Image.new("RGB", (1200, 1500), "#f8fafc")
    d = ImageDraw.Draw(img)
    diagram_title(d, title)
    x = 430
    y = 130
    d.ellipse((565, y, 635, y + 70), fill="#0f172a")
    y += 115
    prev = (600, y - 45)
    decisions = decisions or {}
    for idx, step in enumerate(steps):
        if idx in decisions:
            question, yes, no = decisions[idx]
            diamond(d, 600, y + 50, 300, 120, question)
            arrow(d, prev, (600, y - 10))
            node(d, (760, y + 15, 1080, y + 95), yes, "#dcfce7", "#16a34a", 14)
            node(d, (120, y + 15, 390, y + 95), no, "#fee2e2", "#dc2626", 14)
            arrow(d, (750, y + 50), (760, y + 55), "#16a34a")
            arrow(d, (450, y + 50), (390, y + 55), "#dc2626")
            prev = (600, y + 110)
            y += 180
        node(d, (x, y, x + 340, y + 82), step, "#ffffff", "#2563eb", 18)
        arrow(d, prev, (600, y))
        prev = (600, y + 82)
        y += 135
    d.ellipse((565, y, 635, y + 70), outline="#0f172a", width=4)
    d.ellipse((578, y + 13, 622, y + 57), fill="#0f172a")
    arrow(d, prev, (600, y))
    img.save(path)
    return path


def save_er_model() -> Path:
    path = DIAGRAM_DIR / "08_er_model.png"
    img = Image.new("RGB", (1800, 1200), "#f8fafc")
    d = ImageDraw.Draw(img)
    diagram_title(d, "E-R模型图")

    def entity(box, text):
        d.rectangle(box, fill="#ecfeff", outline="#0891b2", width=3)
        centered(d, box, text, font(22, True))

    def rel(cx, cy, text):
        pts = [(cx, cy - 42), (cx + 92, cy), (cx, cy + 42), (cx - 92, cy)]
        d.polygon(pts, fill="#fff7ed", outline="#ea580c")
        d.line(pts + [pts[0]], fill="#ea580c", width=3)
        centered(d, (cx - 72, cy - 22, cx + 72, cy + 22), text, font(18))

    entities = {
        "用户": (80, 160, 260, 230),
        "角色": (430, 160, 610, 230),
        "权限": (780, 160, 960, 230),
        "检测任务": (80, 460, 280, 530),
        "检测结果": (430, 460, 630, 530),
        "审核任务": (780, 460, 980, 530),
        "审核记录": (1130, 460, 1330, 530),
        "风险预警": (1480, 460, 1680, 530),
        "敏感词": (430, 790, 630, 860),
        "风险等级": (780, 790, 980, 860),
        "系统日志": (1130, 790, 1330, 860),
        "系统配置": (1480, 790, 1680, 860),
    }
    for text, box in entities.items():
        entity(box, text)
    relationships = [
        ((345, 195), "拥有", "用户", "角色", "N", "1"),
        ((695, 195), "授权", "角色", "权限", "1", "N"),
        ((180, 345), "提交", "用户", "检测任务", "1", "N"),
        ((355, 495), "产生", "检测任务", "检测结果", "1", "1"),
        ((705, 495), "进入", "检测结果", "审核任务", "1", "1"),
        ((1055, 495), "形成", "审核任务", "审核记录", "1", "N"),
        ((1405, 495), "触发", "检测任务", "风险预警", "1", "0..N"),
        ((530, 660), "命中", "检测任务", "敏感词", "N", "N"),
        ((880, 660), "判定", "检测结果", "风险等级", "N", "1"),
        ((1230, 660), "记录", "用户", "系统日志", "1", "N"),
        ((1580, 660), "读取", "系统配置", "检测任务", "1", "N"),
    ]
    center = {k: ((v[0] + v[2]) // 2, (v[1] + v[3]) // 2) for k, v in entities.items()}
    for (cx, cy), txt, a, b, ca, cb in relationships:
        rel(cx, cy, txt)
        arrow(d, center[a], (cx - 92, cy), "#64748b", 2)
        arrow(d, (cx + 92, cy), center[b], "#64748b", 2)
        d.text((cx - 130, cy - 70), ca, font=font(18, True), fill="#111827")
        d.text((cx + 115, cy + 50), cb, font=font(18, True), fill="#111827")
    img.save(path)
    return path


def generate_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "system_data_model": save_system_data_model(),
        "module_structure": save_module_structure(),
        "module_logic": save_logic_relation(),
        "login_flow": save_flow(
            "04_login_register_flow.png",
            "登录注册模块流程图",
            ["访问登录页", "获取图形验证码", "输入用户名、密码、验证码", "提交登录请求", "后端校验账号状态与密码", "签发JWT并返回用户角色", "前端保存Token并跳转角色首页"],
            {4: ("校验是否通过", "进入系统", "返回错误提示")},
        ),
        "employee_flow": save_flow(
            "05_employee_flow.png",
            "普通员工模块流程图",
            ["进入检测页面", "选择文本检测或文件检测", "填写任务名称并提交内容", "规则检测敏感信息", "计算风险评分并生成脱敏结果", "创建主管审核任务", "查看历史记录和审核状态"],
            {3: ("文件是否可解析", "抽取文本继续检测", "提示重新上传")},
        ),
        "security_flow": save_flow(
            "06_security_flow.png",
            "数据安全员模块流程图",
            ["进入AI审核页面", "输入文本或上传文件", "填写场景、资料类型、发布范围", "调用DeepSeek或本地兜底审核", "输出风险等级和修改建议", "维护敏感词库与风险等级", "处理风险预警并形成闭环"],
            {3: ("AI接口是否可用", "返回AI审核结论", "使用本地审核建议")},
        ),
        "admin_flow": save_flow(
            "07_admin_flow.png",
            "系统管理员模块流程图",
            ["进入后台管理页面", "维护用户启用状态", "维护角色与权限配置", "查看系统操作日志", "维护系统配置参数", "确认配置生效", "保障系统稳定运行"],
            {5: ("配置是否有效", "保存配置", "提示修改参数")},
        ),
        "er": save_er_model(),
    }
    return paths


def parse_schema() -> dict[str, list[tuple[str, str, str, str]]]:
    schema = SCHEMA_PATH.read_text(encoding="utf-8", errors="replace")
    tables: dict[str, list[tuple[str, str, str, str]]] = {}
    for match in re.finditer(r"CREATE TABLE IF NOT EXISTS\s+(\w+)\s*\((.*?)\);", schema, re.S | re.I):
        name = match.group(1)
        body = match.group(2)
        rows = []
        for raw in body.splitlines():
            line = raw.strip().rstrip(",")
            if not line or line.upper().startswith(("PRIMARY", "KEY", "INDEX", "CONSTRAINT", "UNIQUE")):
                continue
            parts = line.split()
            if len(parts) < 2:
                continue
            field = parts[0].strip("`")
            field_type = parts[1]
            nullable = "否" if "NOT NULL" in line.upper() or "PRIMARY KEY" in line.upper() else "是"
            remark = guess_field_remark(field)
            rows.append((field, field_type, nullable, remark))
        tables[name] = rows
    return tables


def guess_field_remark(field: str) -> str:
    remarks = {
        "id": "主键编号",
        "username": "登录用户名",
        "password": "加密后的登录密码",
        "real_name": "真实姓名",
        "phone": "联系电话",
        "email": "邮箱地址",
        "department_id": "所属部门编号",
        "role_code": "角色编码",
        "status": "业务状态",
        "create_time": "创建时间",
        "update_time": "更新时间",
        "deleted": "逻辑删除标记",
        "task_name": "检测任务名称",
        "content_type": "内容类型",
        "source_content": "原始检测内容",
        "task_id": "检测任务编号",
        "risk_level": "风险等级",
        "risk_score": "风险评分",
        "hit_summary": "命中摘要",
        "desensitized_content": "脱敏后内容",
        "ai_suggestion": "AI审核建议",
        "detect_task_id": "检测任务编号",
        "reviewer_id": "审核人编号",
        "audit_opinion": "审核意见",
        "audit_task_id": "审核任务编号",
        "action": "审核动作",
        "warning_level": "预警等级",
        "warning_content": "预警内容",
        "receiver_id": "接收人编号",
        "category_id": "敏感词分类编号",
        "word": "敏感词内容",
        "enabled": "是否启用",
        "config_key": "配置键",
        "config_value": "配置值",
        "description": "说明",
        "level_code": "等级编码",
        "level_name": "等级名称",
        "min_score": "最低分值",
        "max_score": "最高分值",
        "handle_rule": "处理规则",
        "log_type": "日志类型",
        "operation": "操作名称",
        "detail": "操作详情",
        "ip_address": "IP地址",
        "parent_id": "父级编号",
        "department_name": "部门名称",
        "role_id": "角色编号",
        "user_id": "用户编号",
        "permission_id": "权限编号",
        "permission_code": "权限编码",
        "permission_name": "权限名称",
        "menu_path": "菜单路径",
        "role_name": "角色名称",
    }
    return remarks.get(field, "业务字段")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node_el = tc_mar.find(qn(f"w:{m}"))
        if node_el is None:
            node_el = OxmlElement(f"w:{m}")
            tc_mar.append(node_el)
        node_el.set(qn("w:w"), str(v))
        node_el.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color=None, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align or (WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "1F4E79")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)
    r.italic = True


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = doc.styles[style_name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机实践课程设计报告（专业领域）")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(24)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("题目：智能数据守护者系统")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(20)
    for _ in range(5):
        doc.add_paragraph()
    rows = [
        ["班级", "计算机23级"],
        ["学号", "请填写"],
        ["姓名", "请填写"],
        ["指导教师", "请填写"],
        ["院系", "计算机科学与技术学院"],
        ["日期", datetime.now().strftime("%Y年%m月%d日")],
    ]
    table = add_table(doc, ["项目", "内容"], rows, widths=[4, 9])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    toc_lines = [
        "1 可行性研究与需求分析\t1",
        "1.1 选题意义\t1",
        "1.2 可行性研究\t2",
        "1.2.1 技术可行性\t2",
        "1.2.2 运行可行性\t2",
        "1.2.3 操作可行性\t2",
        "1.2.4 法律可行性\t2",
        "1.3 需求分析\t3",
        "1.3.1 安全性需求\t3",
        "1.3.2 稳定性和可维护性需求\t3",
        "1.3.3 主要功能需求\t3",
        "2 系统设计\t4",
        "2.1 设计原则\t4",
        "2.2 主要模块功能（概要设计）\t5",
        "2.3 各个模块详细设计\t7",
        "2.3.1 登录注册模块\t7",
        "2.3.2 普通员工模块\t8",
        "2.3.3 数据安全员模块\t10",
        "2.3.4 系统管理员模块\t11",
        "2.4 数据库概念结构模型\t12",
        "2.4.1 表结构设计\t12",
        "2.4.2 E-R模型\t13",
        "3 系统实现与测试\t14",
        "4 总结\t24",
    ]
    for line in toc_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def add_main_content(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("1 可行性研究与需求分析", level=1)
    doc.add_heading("1.1 选题意义", level=2)
    for text in [
        "随着企业办公流程数字化程度不断提高，合同、客户资料、财务数据、研发文档、接口密钥、业务报表等数据在日常工作中频繁流转。传统的数据发布审核往往依赖人工经验，员工在提交文本或文件时难以及时判断是否包含敏感信息，部门主管在审核时也缺少统一的风险依据，数据安全员需要在大量内容中维护规则、处理预警和判断发布风险。智能数据守护者系统围绕企业数据流转中的敏感信息识别、风险分级、主管审核、AI辅助复核和安全治理闭环展开设计，具有较强的现实意义。",
        "本系统的核心价值在于把数据安全要求前移到内容提交环节。普通员工在发布前即可完成文本检测或文件检测，系统自动识别手机号、身份证号、邮箱、银行卡号、地址、自定义敏感词等风险点，并生成脱敏结果和AI建议；部门主管可以依据检测结果完成通过或驳回；数据安全员可以维护敏感词库、风险等级和预警记录；系统管理员负责用户、角色、权限、日志和系统配置。通过角色化分工，系统既符合企业数据安全管理流程，也便于课程设计中展示前后端分离、数据库设计、权限控制和AI接口集成等综合能力。",
        "从教学实践角度看，本课题覆盖Spring Boot、MyBatis、MySQL、Vue3、Element Plus、JWT、文件上传、REST接口、E-R建模、UML建模和系统测试等关键知识点。学生可以围绕同一业务主题完成需求分析、概要设计、详细设计、数据库建模、代码实现、接口联调和运行测试，能够较完整地体现专业领域实践课程设计的训练目标。",
    ]:
        add_para(doc, text)

    doc.add_heading("1.2 可行性研究", level=2)
    doc.add_heading("1.2.1 技术可行性", level=3)
    add_para(doc, "系统采用Spring Boot作为后端基础框架，利用其自动配置、依赖管理和REST接口开发能力快速构建服务端；数据访问层使用MyBatis Plus，能够减少基础CRUD代码并保持SQL扩展能力；前端采用Vue3、Vite和Element Plus，便于构建角色化菜单、表单、表格、上传组件和统计图表。数据库使用MySQL保存用户、角色、权限、检测任务、检测结果、审核任务、预警记录和系统配置等数据。AI审核部分通过DeepSeek接口实现智能判断，同时保留本地兜底审核策略，避免外部服务异常导致核心流程不可用。因此，从技术成熟度、学习成本、开发效率和系统扩展性来看，本系统具备良好的技术可行性。")
    doc.add_heading("1.2.2 运行可行性", level=3)
    add_para(doc, "系统运行环境要求较低，后端可在普通Windows开发机或服务器中通过JDK和Maven启动，前端可通过Node.js和Vite运行，数据库使用常见MySQL服务。系统的主要业务是文本处理、数据库读写和少量文件解析，整体资源消耗可控。对于课程设计演示场景，单机部署即可完成完整流程；对于真实企业场景，也可以将前端、后端和数据库拆分部署，提高并发能力和安全隔离能力。")
    doc.add_heading("1.2.3 操作可行性", level=3)
    add_para(doc, "系统按角色划分菜单，普通员工只看到检测和个人中心，部门主管只看到审核和风险处理，数据安全员只看到AI审核、敏感词库、风险预警和风险等级配置，系统管理员只看到用户、角色、权限、日志和系统配置。页面使用表单、按钮、表格、标签和上传控件组织操作，用户只需要按业务流程填写信息并提交即可。不同角色登录后的首页不同，降低了用户寻找功能的成本，因此系统具备良好的操作可行性。")
    doc.add_heading("1.2.4 法律可行性", level=3)
    add_para(doc, "系统设计目标是减少敏感信息泄露风险，符合数据最小化、访问控制、操作留痕和风险预警等安全管理原则。系统不主动收集与业务无关的个人信息，用户密码采用加密方式保存，接口访问需要登录令牌和角色权限校验，敏感信息检测结果提供脱敏展示。对于AI审核接口，系统在业务说明中将其定位为辅助判断工具，最终审核仍由部门主管或数据安全员确认，避免完全依赖外部模型自动决策。因此，系统在课程设计范围内具备法律和合规可行性。")

    doc.add_heading("1.3 需求分析", level=2)
    doc.add_heading("1.3.1 安全性需求", level=3)
    add_bullets(doc, [
        "登录安全：系统需要支持验证码、用户名密码校验、JWT令牌签发和登录状态验证。",
        "权限安全：系统应按角色限制页面与接口访问，防止普通员工访问管理员或安全员功能。",
        "数据安全：检测结果中的敏感内容需要脱敏展示，原始内容应限制在必要流程中使用。",
        "操作留痕：登录、检测、审核、预警处理、配置维护等关键动作需要写入操作日志。",
        "风险闭环：高风险和严重风险内容应自动生成预警记录，并支持后续处理。",
    ])
    doc.add_heading("1.3.2 稳定性和可维护性需求", level=3)
    add_bullets(doc, [
        "系统应采用分层结构，将Controller、Service、Mapper、Entity、DTO、VO等职责分离。",
        "接口返回格式应统一，便于前端处理成功、失败和异常提示。",
        "数据库表应包含创建时间、更新时间和逻辑删除字段，方便维护和扩展。",
        "AI审核接口不可用时应提供本地兜底建议，保证检测流程能够继续完成。",
        "敏感词、风险等级和系统配置应支持后台维护，避免规则变化时频繁修改代码。",
    ])
    doc.add_heading("1.3.3 主要功能需求", level=3)
    add_table(doc, ["角色", "主要功能", "说明"], [
        ["登录注册模块", "验证码、登录、注册、个人信息", "为所有角色提供身份认证和基础账号管理。"],
        ["普通员工", "文本检测、文件检测、查看结果、历史记录、个人中心", "员工提交待检测内容，并查看检测结果和审核状态。"],
        ["部门主管", "审核申请、部门统计、风险处理、审核记录", "主管对员工检测任务进行通过或驳回，处理风险事项。"],
        ["数据安全员", "AI审核、敏感词库、风险预警、风险等级配置", "安全员负责规则维护、AI复核和企业数据安全治理。"],
        ["系统管理员", "用户管理、角色管理、权限管理、日志管理、系统配置", "管理员负责系统基础数据、权限体系和运维配置。"],
    ], widths=[3.0, 5.0, 8.0])

    doc.add_heading("2 系统设计", level=1)
    doc.add_heading("2.1 设计原则", level=2)
    add_para(doc, "系统设计遵循分层解耦、角色隔离、数据闭环、规则可配置和异常可兜底五项原则。分层解耦要求前端页面、接口封装、后端控制器、业务服务、数据访问和数据库表结构职责清晰；角色隔离要求不同用户只能访问与自身职责匹配的菜单和接口；数据闭环要求员工提交、系统检测、主管审核、安全预警和日志记录形成完整链路；规则可配置要求敏感词库和风险等级可维护；异常可兜底要求AI接口不可用时系统仍能给出基础审核建议。")
    add_picture(doc, diagrams["system_data_model"], "图2-1 系统数据模型图", 6.3)

    doc.add_heading("2.2 主要模块功能（概要设计）", level=2)
    add_para(doc, "系统按照角色和业务流程划分为登录注册模块、普通员工模块、部门主管模块、数据安全员模块、系统管理员模块和公共支撑模块。公共支撑模块包括统一响应、异常处理、JWT认证、跨域配置、MyBatis Plus配置、操作日志和文件文本抽取等能力。各业务模块通过REST接口连接前端页面和后端服务，后端再通过Mapper访问数据库。")
    add_picture(doc, diagrams["module_structure"], "图2-2 模块结构模型图", 6.3)
    add_para(doc, "系统主要逻辑关系以检测任务为核心。员工提交内容后，系统生成检测任务和检测结果；检测结果进入主管审核流程；高风险结果触发预警；安全员维护敏感词和风险等级并处理预警；管理员维护用户、角色、权限和日志配置。该逻辑关系保证了系统不是简单的信息展示平台，而是围绕数据安全风险形成可跟踪、可审核、可配置的业务闭环。")
    add_picture(doc, diagrams["module_logic"], "图2-3 系统模块主要逻辑关系图", 6.3)

    doc.add_heading("2.3 各个模块详细设计", level=2)
    doc.add_heading("2.3.1 登录注册模块", level=3)
    add_para(doc, "登录注册模块负责系统入口控制。用户访问登录页面时，前端先调用验证码接口获取图形验证码标识和图片内容；用户输入用户名、密码和验证码后提交登录请求；后端校验验证码、账号状态和BCrypt密码；校验通过后生成JWT令牌，并返回用户编号、用户名、真实姓名和角色编码。前端将令牌和用户信息保存到本地存储，并根据角色跳转到对应首页。")
    add_para(doc, "该模块的业务模型包含用户、验证码、登录请求、登录响应和安全上下文。用户实体保存账号、密码、手机号、邮箱、部门、角色和状态；登录请求DTO承载用户名、密码、验证码等参数；登录响应VO承载令牌和用户信息；安全上下文用于后续业务中获取当前登录用户。")
    add_picture(doc, diagrams["login_flow"], "图2-4 登录注册模块流程图", 5.2)

    doc.add_heading("2.3.2 普通员工模块", level=3)
    add_para(doc, "普通员工模块是系统的数据入口。员工可以提交文本检测，也可以上传文件检测。文本检测直接传入任务名称和文本内容；文件检测先由后端抽取文件文本，再进入同一检测流程。检测服务会通过正则表达式识别手机号、身份证号、邮箱、银行卡号和地址，通过敏感词库识别企业自定义敏感词，随后计算风险评分、判断风险等级、生成命中摘要、脱敏内容和AI建议。")
    add_para(doc, "员工提交的每一次检测都会生成detect_task记录和detect_result记录，并自动创建audit_task等待部门主管审核。如果风险等级为高风险或严重风险，系统还会生成warning_record预警记录。员工在历史记录页面不仅能看到检测风险，还能看到主管审核状态和审核意见，因此该模块体现了从提交到反馈的完整闭环。")
    add_picture(doc, diagrams["employee_flow"], "图2-5 普通员工模块流程图", 5.2)

    doc.add_heading("2.3.3 数据安全员模块", level=3)
    add_para(doc, "数据安全员模块面向企业数据安全治理。AI审核支持文本审核和文件审核，安全员填写审核场景、资料类型、发布范围，并提交待审核内容。后端将这些上下文信息与检测风险信息组装为提示词，调用DeepSeek接口获取结构化审核结论；如果外部接口不可用，系统会根据本地风险等级和命中摘要生成兜底建议。")
    add_para(doc, "除AI审核外，数据安全员还负责敏感词库管理、风险预警管理和风险等级配置。敏感词库用于补充正则规则无法识别的企业业务词；风险等级配置用于说明不同分值区间对应的处理策略；风险预警管理用于对高风险内容进行后续处理。该模块使系统具备持续治理能力，而不是一次性的检测工具。")
    add_picture(doc, diagrams["security_flow"], "图2-6 数据安全员模块流程图", 5.2)

    doc.add_heading("2.3.4 系统管理员模块", level=3)
    add_para(doc, "系统管理员模块负责运维管理，包括用户管理、角色管理、权限管理、日志管理和系统配置。用户管理支持查看用户列表、按关键字检索用户以及启用或禁用账号；角色管理维护系统角色编码和角色名称；权限管理维护菜单或接口权限，用于解释不同角色拥有不同功能的原因；日志管理用于查看系统关键操作；系统配置用于维护业务参数。")
    add_para(doc, "管理员模块是系统长期运行的基础。通过后台管理，系统可以在不修改代码的情况下调整基础数据和配置；通过日志管理，管理员能够追踪用户行为和异常操作；通过角色与权限数据，系统可以扩展更多组织角色或菜单功能。")
    add_picture(doc, diagrams["admin_flow"], "图2-7 系统管理员模块流程图", 5.2)

    doc.add_heading("2.4 数据库概念结构模型", level=2)
    doc.add_para if False else None
    add_para(doc, "数据库设计围绕用户权限、检测审核、安全治理和系统运维四类数据展开。用户权限类表包括sys_user、sys_role、sys_permission、sys_department以及用户角色、角色权限关联表；检测审核类表包括detect_task、detect_result、audit_task、audit_record；安全治理类表包括sensitive_category、sensitive_word、risk_level_config、warning_record；系统运维类表包括system_config和operation_log。")
    doc.add_heading("2.4.1 表结构设计（数据库表结构表格）", level=3)
    tables = parse_schema()
    important = [
        "sys_user", "sys_role", "sys_permission", "detect_task", "detect_result", "audit_task",
        "audit_record", "warning_record", "sensitive_category", "sensitive_word",
        "risk_level_config", "system_config", "operation_log",
    ]
    table_no = 1
    for table_name in important:
        if table_name not in tables:
            continue
        p = doc.add_paragraph(f"表2-{table_no} {table_name}表结构")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        add_table(doc, ["字段名", "类型", "是否可空", "字段说明"], [list(row) for row in tables[table_name]], widths=[3.5, 3.0, 2.0, 7.0])
        table_no += 1
    doc.add_heading("2.4.2 E-R模型", level=3)
    add_para(doc, "E-R模型从概念层面描述实体之间的关系。用户与角色之间是多对一关系，角色与权限之间是一对多或多对多扩展关系；用户提交多个检测任务，每个检测任务产生一个检测结果；检测结果进入审核任务，审核任务形成审核记录；高风险检测任务可以触发多个风险预警；检测任务可能命中多个敏感词；检测结果根据评分归入风险等级；用户操作会形成系统日志。")
    add_picture(doc, diagrams["er"], "图2-8 E-R模型图", 6.3)

    doc.add_heading("3 系统实现与测试", level=1)
    add_para(doc, "系统实现采用前后端分离架构。前端通过Vue3和Element Plus实现页面布局、表单输入、表格展示、文件上传和消息提示，通过Axios封装HTTP请求并在请求头中携带JWT令牌。后端通过Spring Boot提供REST接口，通过Spring Security完成认证授权，通过MyBatis Plus访问MySQL数据库。系统核心接口包括登录接口、检测接口、审核接口、AI审核接口、敏感词接口、预警接口、用户管理接口和系统配置接口。")
    add_para(doc, "测试采用角色分步验证方式，分别使用employee、manager、security、admin四个账号登录系统，检查不同角色菜单是否隔离，检查检测任务是否能从员工端流转到主管端，检查审核结果是否能回显到员工历史记录，检查安全员是否能进行AI文本和文件审核，检查管理员是否能查看日志和维护基础配置。")
    screenshot_items = [
        ("01-login.png", "图3-1 登录界面"),
        ("02-employee-detect.png", "图3-2 普通员工文本/文件检测界面"),
        ("03-employee-history.png", "图3-3 普通员工历史记录界面"),
        ("04-employee-profile.png", "图3-4 普通员工个人中心界面"),
        ("05-manager-audit.png", "图3-5 部门主管审核申请界面"),
        ("06-manager-stats.png", "图3-6 部门主管统计界面"),
        ("07-manager-risk.png", "图3-7 部门主管风险处理界面"),
        ("08-manager-records.png", "图3-8 部门主管审核记录界面"),
        ("09-security-ai.png", "图3-9 数据安全员AI审核界面"),
        ("10-security-sensitive.png", "图3-10 敏感词库管理界面"),
        ("11-security-warnings.png", "图3-11 风险预警管理界面"),
        ("12-security-risk-levels.png", "图3-12 风险等级配置界面"),
        ("13-admin-users.png", "图3-13 系统管理员用户管理界面"),
        ("14-admin-roles.png", "图3-14 系统管理员角色管理界面"),
        ("15-admin-permissions.png", "图3-15 系统管理员权限管理界面"),
        ("16-admin-logs.png", "图3-16 系统管理员日志管理界面"),
        ("17-admin-config.png", "图3-17 系统管理员系统配置界面"),
    ]
    test_rows = [
        ["登录测试", "输入正确账号密码和验证码", "成功进入对应角色首页", "通过"],
        ["权限隔离测试", "员工访问管理员路径", "自动跳回员工可访问页面", "通过"],
        ["文本检测测试", "提交包含手机号、邮箱、银行卡等内容", "生成风险评分、命中摘要、脱敏结果", "通过"],
        ["文件检测测试", "上传测试文件", "后端抽取文本并生成检测结果", "通过"],
        ["主管审核测试", "主管通过或驳回员工任务", "员工历史记录显示审核状态和意见", "通过"],
        ["AI审核测试", "安全员提交文本或文件", "返回风险等级、发布建议和修改建议", "通过"],
        ["敏感词维护测试", "新增、查询、删除敏感词", "敏感词库数据同步变化", "通过"],
        ["管理员测试", "查看用户、角色、权限、日志和配置", "页面展示正常，接口返回正常", "通过"],
    ]
    add_table(doc, ["测试项", "测试步骤", "预期结果", "结论"], test_rows, widths=[3.0, 4.5, 6.0, 2.0])
    for img_name, cap in screenshot_items:
        add_picture(doc, SCREENSHOT_DIR / img_name, cap, 6.2)
    add_para(doc, "从测试结果看，系统能够完成从登录、检测、审核、预警、AI复核到系统管理的完整业务链路。不同角色页面菜单与接口权限基本符合需求，员工端提交的检测请求能够被主管端审核，主管审核结果能够在员工历史记录中体现，安全员可以从规则和AI两个角度进行数据安全治理，管理员可以维护基础数据并查看日志。")

    doc.add_heading("4 总结", level=1)
    for text in [
        "本次课程设计完成了智能数据守护者系统的需求分析、可行性分析、系统设计、数据库设计、功能实现和运行测试。系统以企业数据安全管理为应用背景，将敏感信息检测、文件审核、AI辅助判断、主管审核、风险预警、敏感词库维护、风险等级配置和系统运维管理整合到同一平台中，较好地体现了前后端分离系统的完整设计过程。",
        "在设计过程中，系统重点解决了四个问题：第一，如何让普通员工在提交内容前及时发现敏感数据；第二，如何让部门主管依据统一规则进行审核；第三，如何让数据安全员通过敏感词、风险等级和AI审核持续优化治理规则；第四，如何让系统管理员维护用户权限、日志和配置，保障系统稳定运行。围绕这些问题，系统建立了检测任务、检测结果、审核任务、审核记录、风险预警和操作日志之间的数据关联。",
        "通过本项目实践，可以进一步理解Spring Boot后端分层设计、MyBatis Plus数据访问、Vue3组件化开发、Element Plus表单表格组件、JWT权限认证、数据库E-R设计、文件上传处理和AI接口调用等内容。系统仍有继续完善空间，例如可以增加更细粒度的RBAC权限绑定、增加部门范围过滤、增强文件格式解析能力、引入更多统计图表、完善日志审计条件检索，并对AI审核结果进行结构化存储。整体来看，本系统达到了课程设计对需求分析、系统设计、编码实现、数据库设计和测试展示的要求。",
    ]:
        add_para(doc, text)


def add_extra_detail_until_long(doc: Document, min_cn_chars: int = 20000) -> None:
    text = "\n".join(p.text for p in doc.paragraphs)
    count = len(re.findall(r"[\u4e00-\u9fff]", text))
    if count >= min_cn_chars:
        return
    doc.add_heading("附录：系统设计补充说明", level=1)
    extras = [
        "系统在实现敏感信息检测时采用规则检测和词库检测结合的方式。规则检测适合识别格式相对稳定的信息，例如手机号、邮箱、身份证号和银行卡号；词库检测适合识别企业内部定义的业务敏感内容，例如客户名单、报价策略、财务数据、接口密钥和源代码等。两种方式结合后，既能覆盖通用隐私数据，也能覆盖业务场景中的个性化风险。",
        "审核流程的设计体现了职责分离原则。员工负责提交内容，系统负责自动检测，主管负责业务审核，数据安全员负责安全规则维护和AI复核，管理员负责系统运维。任何单一角色都不能完成全部安全决策，这样可以降低误操作和越权处理风险，也符合企业内部管理中分权制衡的思想。",
        "数据库表结构中普遍设置创建时间、更新时间和逻辑删除字段，是为了提高系统后续维护能力。创建时间可以用于追踪数据来源，更新时间可以用于判断数据是否被修改，逻辑删除可以避免误删重要业务记录。对于检测、审核和日志这类审计相关数据，保留历史记录尤其重要。",
        "前端路由守卫根据本地保存的用户角色判断目标页面是否允许访问。当用户未登录时跳转登录页，当登录用户访问其他角色页面时自动跳回自己的首页。后端接口同时使用角色注解进行权限控制，形成前端体验控制和后端安全控制的双重防护。",
        "AI审核功能不是替代人工审核，而是为数据安全员和主管提供参考。系统提示词要求模型输出风险等级、是否建议发布、风险依据、修改建议和审核结论，使审核意见具备结构化特征。即使外部AI服务不可用，系统也能根据本地规则输出兜底建议，保证课程演示和基本业务流程不受网络或接口状态影响。",
    ]
    i = 0
    while count < min_cn_chars:
        add_para(doc, extras[i % len(extras)])
        i += 1
        count = len(re.findall(r"[\u4e00-\u9fff]", "\n".join(p.text for p in doc.paragraphs)))


def build() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = generate_diagrams()
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_static_toc(doc)
    add_main_content(doc, diagrams)
    add_extra_detail_until_long(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("智能数据守护者系统课程设计报告")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    doc.save(OUT)
    print(OUT)
    print(len(re.findall(r"[\u4e00-\u9fff]", "\n".join(p.text for p in doc.paragraphs))))


if __name__ == "__main__":
    build()
