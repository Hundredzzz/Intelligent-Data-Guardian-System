from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from build_course_report import (
    REPORT_DIR,
    DIAGRAM_DIR,
    SCREENSHOT_DIR,
    build_diagrams,
    set_doc_styles,
    set_cell_text,
    add_table,
    add_image,
    paragraph,
    bullet,
)


OUT = REPORT_DIR / "智能数据守护者系统课程设计报告-模板版.docx"


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("计算机实践课程设计报告（专业领域）")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(24)

    doc.add_paragraph()
    rows = [
        ("题    目", "智能数据守护者系统的设计与实现"),
        ("班    级", "计算机23-*班"),
        ("学    号", "请填写"),
        ("姓    名", "请填写"),
        ("指导教师", "请填写"),
        ("系 主 任", "李成严"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        set_cell_text(table.rows[i].cells[0], row[0], True)
        set_cell_text(table.rows[i].cells[1], row[1])
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("计算机科学与技术学院")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年6月")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(14)
    doc.add_page_break()


def add_catalog(doc: Document):
    doc.add_heading("目录", level=1)
    for item in [
        "1 可行性研究与需求分析",
        "1.1 选题意义",
        "1.2 可行性研究",
        "1.2.1 技术可行性",
        "1.2.2 运行可行性",
        "1.2.3 操作可行性",
        "1.2.4 法律可行性",
        "1.3 需求分析",
        "2 系统设计",
        "2.1 设计原则",
        "2.2 主要模块功能模型",
        "2.3 客户端与服务器端的交互设计模型",
        "2.4 数据库概念结构模型",
        "3 系统实现与测试",
        "3.1 后端系统实现",
        "3.2 前端系统实现",
        "3.3 DeepSeek AI审核实现",
        "3.4 系统运行样例",
        "3.5 功能测试",
        "4 总结",
    ]:
        paragraph(doc, item)
    doc.add_page_break()


def build_report():
    REPORT_DIR.mkdir(exist_ok=True)
    build_diagrams()
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc)
    add_catalog(doc)

    doc.add_heading("1 可行性研究与需求分析", level=1)
    doc.add_heading("1.1 选题意义", level=2)
    paragraph(doc, "随着企业数字化办公和互联网业务的发展，客户资料、财务数据、合同内容、技术资料等信息在内部传输、共享和对外发布过程中频繁流转。传统人工审核方式存在效率低、准确率不足、审核记录不完整等问题，容易造成敏感信息泄露。智能数据守护者系统通过规则识别、敏感词库、DeepSeek AI审核、风险预警和多角色协同审核机制，对数据发布前进行自动化检查与复核，能够提高企业数据安全管理能力。")
    paragraph(doc, "本课题结合 Spring Boot、MyBatis Plus、Vue3、MySQL、JWT、DeepSeek API 等技术，实现一个具有实际业务闭环的数据安全审核平台。系统不仅满足课程设计对需求分析、数据库设计、UML建模、编码实现和测试展示的要求，也具有较好的扩展价值。")

    doc.add_heading("1.2 可行性研究", level=2)
    doc.add_heading("1.2.1 技术可行性", level=3)
    paragraph(doc, "系统采用成熟的前后端分离架构。后端 Spring Boot 3 负责 REST 接口、业务流程、安全认证和数据访问；MyBatis Plus 简化数据库 CRUD；MySQL 8.0 存储用户、角色、权限、检测任务、审核记录、预警记录和日志；前端 Vue3、Vite、Element Plus 实现多角色管理界面；DeepSeek API 用于 AI 辅助审核。上述技术生态成熟、资料丰富、兼容性良好，因此技术上可行。")
    doc.add_heading("1.2.2 运行可行性", level=3)
    paragraph(doc, "系统运行环境主要包括 JDK 17、Maven、Node.js、MySQL 8.0 和浏览器。项目可在普通 Windows 开发机上运行，后端默认端口 8080，前端默认端口 5173，部署与调试成本较低。")
    doc.add_heading("1.2.3 操作可行性", level=3)
    paragraph(doc, "系统按角色划分菜单，普通员工、部门主管、数据安全员、系统管理员登录后只看到本角色功能。页面以表单、表格、上传控件、弹窗为主，操作路径清晰，适合课程演示和普通用户使用。")
    doc.add_heading("1.2.4 法律可行性", level=3)
    paragraph(doc, "项目使用开源框架和课程设计自研代码，DeepSeek API 用于合法的文本审核场景。系统测试样本为模拟数据，不涉及真实个人隐私或商业秘密，因此法律上可行。")

    doc.add_heading("1.3 需求分析", level=2)
    add_table(doc, ["角色", "主要职责", "功能需求"], [
        ["普通员工", "提交检测内容", "文本检测、文件检测、查看检测结果、查看历史记录、个人中心"],
        ["部门主管", "审核员工内容", "审核申请、部门统计、风险处理、审核记录"],
        ["数据安全员", "企业数据安全管理", "AI审核、敏感词库管理、风险预警管理、风险等级配置"],
        ["系统管理员", "系统运维", "用户管理、角色管理、权限管理、日志管理、系统配置"],
    ])
    for item in [
        "安全性需求：采用 JWT 鉴权、BCrypt 密码加密、角色权限控制和操作日志审计。",
        "稳定性和可维护性需求：采用分层架构，控制器、服务、数据访问和实体类职责清晰。",
        "主要功能需求：覆盖敏感信息检测、AI审核、主管审核、风险预警、敏感词维护、统计分析和系统运维。",
    ]:
        bullet(doc, item)

    doc.add_heading("2 系统设计", level=1)
    doc.add_heading("2.1 设计原则", level=2)
    for item in [
        "安全优先：敏感信息识别、脱敏展示、风险预警和日志审计贯穿业务流程。",
        "角色分离：不同角色只访问自身职责范围内的功能，前端路由和后端接口共同限制。",
        "模块化设计：认证、检测、AI审核、审核管理、预警、敏感词、权限和日志模块相互独立。",
        "可扩展性：预留 OCR、多模型接入、动态角色权限分配和安全报告生成等扩展方向。",
    ]:
        bullet(doc, item)

    doc.add_heading("2.2 主要模块功能模型", level=2)
    paragraph(doc, "系统主要模块包括用户认证模块、敏感数据检测模块、AI审核模块、审核管理模块、风险预警模块、敏感词管理模块、RBAC权限模块、系统日志模块。各模块通过统一 REST API 与前端交互。")
    add_image(doc, DIAGRAM_DIR / "uml_use_case.png", "图 1 系统 UML 用例图")
    add_image(doc, DIAGRAM_DIR / "uml_class.png", "图 2 系统分层类图")

    doc.add_heading("2.3 客户端与服务器端的交互设计模型", level=2)
    paragraph(doc, "客户端通过 Axios 调用后端接口，后端首先由 JWT 过滤器解析登录用户，再根据角色进行权限校验。检测任务提交后，系统完成规则检测、AI审核建议、结果入库、预警生成和审核任务生成。主管审核后，员工端可查看通过或驳回结果，管理员端可查看操作日志。")
    add_image(doc, DIAGRAM_DIR / "workflow.png", "图 3 系统业务流程图")

    doc.add_heading("2.4 数据库概念结构模型", level=2)
    paragraph(doc, "数据库围绕用户、角色、权限、检测任务、检测结果、审核任务、预警记录、敏感词、风险等级配置和操作日志进行建模。")
    add_image(doc, DIAGRAM_DIR / "er_diagram.png", "图 4 数据库 E-R 图")
    add_table(doc, ["数据表", "作用"], [
        ["sys_user", "保存用户账号、密码、部门和角色"],
        ["sys_role / sys_permission", "保存角色与权限功能点"],
        ["detect_task / detect_result", "保存检测任务和检测结果"],
        ["audit_task / audit_record", "保存主管审核任务和审核记录"],
        ["warning_record", "保存高风险和严重风险预警"],
        ["sensitive_word / sensitive_category", "保存敏感词和分类"],
        ["operation_log", "保存登录、检测、审核、配置等操作日志"],
        ["risk_level_config", "保存风险等级阈值和处理规则"],
    ])

    doc.add_heading("3 系统实现与测试", level=1)
    doc.add_heading("3.1 后端系统实现", level=2)
    paragraph(doc, "后端工程位于 backend 目录，采用 Spring Boot 3 和 MyBatis Plus。主要包包括 controller、service、mapper、entity、dto、vo、config、security、common。系统通过 JwtAuthenticationFilter 解析令牌，通过 @PreAuthorize 限制员工、主管、安全员和管理员接口。")
    doc.add_heading("3.2 前端系统实现", level=2)
    paragraph(doc, "前端工程位于 frontend 目录，采用 Vue3、Vite 和 Element Plus。系统根据登录用户 roleCode 动态生成菜单，并通过路由守卫阻止越权访问。页面覆盖文本/文件检测、历史记录、审核申请、风险处理、AI审核、敏感词库、用户管理、角色管理、权限管理、日志管理和系统配置。")
    doc.add_heading("3.3 DeepSeek AI审核实现", level=2)
    paragraph(doc, "AI审核模块支持文本审核和文件审核。数据安全员选择审核场景、资料类型和发布范围后，系统将内容提交给 DeepSeek API，返回风险等级、是否建议发布、风险依据、修改建议和审核结论。检测流程中的 AI建议也复用该服务。")
    doc.add_heading("3.4 系统运行样例", level=2)
    screenshots = [
        ("01-login.png", "图 5 登录与验证码页面"),
        ("02-employee-detect.png", "图 6 普通员工文本/文件检测页面"),
        ("03-employee-history.png", "图 7 普通员工检测结果与历史记录页面"),
        ("04-employee-profile.png", "图 8 普通员工个人中心页面"),
        ("05-manager-audit.png", "图 9 部门主管审核申请页面"),
        ("06-manager-stats.png", "图 10 部门主管部门统计页面"),
        ("07-manager-risk.png", "图 11 部门主管风险处理页面"),
        ("08-manager-records.png", "图 12 部门主管审核记录页面"),
        ("09-security-ai.png", "图 13 数据安全员 AI审核页面"),
        ("10-security-sensitive.png", "图 14 数据安全员敏感词库管理页面"),
        ("11-security-warnings.png", "图 15 数据安全员风险预警管理页面"),
        ("12-security-risk-levels.png", "图 16 数据安全员风险等级配置页面"),
        ("13-admin-users.png", "图 17 系统管理员用户管理页面"),
        ("14-admin-roles.png", "图 18 系统管理员角色管理页面"),
        ("15-admin-permissions.png", "图 19 系统管理员权限管理页面"),
        ("16-admin-logs.png", "图 20 系统管理员日志管理页面"),
        ("17-admin-config.png", "图 21 系统管理员系统配置页面"),
    ]
    for file, caption in screenshots:
        add_image(doc, SCREENSHOT_DIR / file, caption, width=6.2)

    doc.add_heading("3.5 功能测试", level=2)
    add_table(doc, ["测试编号", "测试功能", "测试账号", "测试内容", "预期结果"], [
        ["T01", "登录认证", "admin", "输入用户名、密码、验证码", "登录成功并进入管理员功能"],
        ["T02", "文本检测", "employee", "输入包含手机号和敏感词的文本", "生成风险结果和审核任务"],
        ["T03", "文件检测", "employee", "上传 txt/pdf/doc/docx 文件", "提取内容并完成检测"],
        ["T04", "主管审核", "manager", "查看检测详情并通过或驳回", "员工历史显示审核状态和意见"],
        ["T05", "AI审核", "security", "输入文本或上传文件", "DeepSeek 返回结构化审核建议"],
        ["T06", "敏感词维护", "security", "新增或修改敏感词", "后续检测可命中新增词"],
        ["T07", "预警处理", "manager/security", "处理高风险预警", "预警状态更新为已处理"],
        ["T08", "日志管理", "admin", "查看操作日志", "显示登录、检测、审核等操作记录"],
    ])
    paragraph(doc, "测试结果表明，系统功能流程完整，能够支持员工提交、系统检测、AI审核、主管复核、安全员管理和管理员审计的完整闭环。")

    doc.add_heading("4 总结", level=1)
    paragraph(doc, "本项目完成了智能数据守护者系统的需求分析、系统设计、数据库设计、编码实现和功能测试。系统采用 Spring Boot、MyBatis Plus、Vue3、MySQL、JWT 和 DeepSeek API 等技术，具备多角色协同、敏感信息识别、AI辅助审核、风险预警、权限管理和日志审计能力。")
    paragraph(doc, "通过本次课程设计，进一步掌握了软件工程文档编写、UML建模、E-R建模、前后端分离开发、RBAC权限控制、文件解析和 AI 服务集成等实践技能。后续可继续扩展 OCR 图片检测、多模型审核、动态角色权限分配和自动安全报告生成功能。")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
