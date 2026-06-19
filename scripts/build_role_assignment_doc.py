from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "智能数据守护者系统四人分工与代码说明文档.docx"


def rel(path: str) -> str:
    return path.replace("/", "\\")


ROLE_DOCS = [
    {
        "member": "成员一",
        "role_name": "普通员工",
        "username": "employee",
        "password": "123456",
        "role_code": "EMPLOYEE",
        "home": "/detect",
        "responsibility": "负责数据检测入口，提交文本或文件检测内容，查看检测结果、审核状态和个人信息。",
        "functions": [
            "文本检测：输入任务名称和文本内容，提交敏感信息检测。",
            "文件检测：上传文件并由后端抽取文本后执行检测。",
            "查看检测结果：查看风险等级、风险评分、命中内容、脱敏结果和 AI 建议。",
            "查看历史记录：查看自己提交过的检测任务，以及部门主管审核通过或驳回状态。",
            "个人中心：查看并维护账号基本资料，支持修改密码。",
        ],
        "frontend": [
            "frontend/src/views/DetectionView.vue",
            "frontend/src/views/HistoryView.vue",
            "frontend/src/views/ProfileView.vue",
            "frontend/src/api/modules.js",
            "frontend/src/router/index.js",
            "frontend/src/views/LayoutView.vue",
        ],
        "backend": [
            "backend/src/main/java/com/guardian/controller/DetectionController.java",
            "backend/src/main/java/com/guardian/controller/AuthController.java",
            "backend/src/main/java/com/guardian/service/DetectionService.java",
            "backend/src/main/java/com/guardian/service/FileTextExtractor.java",
            "backend/src/main/java/com/guardian/service/AuthService.java",
            "backend/src/main/java/com/guardian/dto/DetectRequest.java",
            "backend/src/main/java/com/guardian/dto/ProfileRequest.java",
            "backend/src/main/java/com/guardian/dto/ChangePasswordRequest.java",
            "backend/src/main/java/com/guardian/entity/DetectTask.java",
            "backend/src/main/java/com/guardian/entity/DetectResult.java",
            "backend/src/main/java/com/guardian/entity/AuditTask.java",
            "backend/src/main/java/com/guardian/entity/WarningRecord.java",
            "backend/src/main/java/com/guardian/mapper/DetectTaskMapper.java",
            "backend/src/main/java/com/guardian/mapper/DetectResultMapper.java",
            "backend/src/main/java/com/guardian/mapper/AuditTaskMapper.java",
            "backend/src/main/java/com/guardian/vo/DetectionResultVO.java",
            "backend/src/main/java/com/guardian/vo/DetectionHistoryVO.java",
            "backend/src/main/java/com/guardian/vo/DetectionHitVO.java",
        ],
        "tables": ["detect_task", "detect_result", "audit_task", "warning_record", "sys_user"],
        "apis": ["/detect/text", "/detect/file", "/detect/history", "/auth/profile", "/auth/password"],
    },
    {
        "member": "成员二",
        "role_name": "部门主管",
        "username": "manager",
        "password": "123456",
        "role_code": "MANAGER",
        "home": "/audit",
        "responsibility": "负责审核员工提交的检测任务，查看部门统计，对风险预警进行处理并留存审核记录。",
        "functions": [
            "审核申请：查看员工提交的待审核检测任务，结合风险等级、命中摘要和脱敏内容进行判断。",
            "通过或驳回：填写审核意见，将审核状态同步回员工历史记录。",
            "部门统计：查看检测任务数量、风险分布等统计信息。",
            "风险处理：处理高风险或严重风险预警，形成风险闭环。",
            "审核记录：查看历史审核动作、审核意见和审核时间。",
        ],
        "frontend": [
            "frontend/src/views/AuditView.vue",
            "frontend/src/views/ManagerStatsView.vue",
            "frontend/src/views/WarningView.vue",
            "frontend/src/views/AuditRecordsView.vue",
            "frontend/src/api/modules.js",
            "frontend/src/router/index.js",
            "frontend/src/views/LayoutView.vue",
        ],
        "backend": [
            "backend/src/main/java/com/guardian/controller/AuditController.java",
            "backend/src/main/java/com/guardian/controller/DashboardController.java",
            "backend/src/main/java/com/guardian/controller/WarningController.java",
            "backend/src/main/java/com/guardian/service/AuditService.java",
            "backend/src/main/java/com/guardian/service/DashboardService.java",
            "backend/src/main/java/com/guardian/service/WarningService.java",
            "backend/src/main/java/com/guardian/dto/AuditRequest.java",
            "backend/src/main/java/com/guardian/entity/AuditTask.java",
            "backend/src/main/java/com/guardian/entity/AuditRecord.java",
            "backend/src/main/java/com/guardian/entity/DetectTask.java",
            "backend/src/main/java/com/guardian/entity/DetectResult.java",
            "backend/src/main/java/com/guardian/entity/WarningRecord.java",
            "backend/src/main/java/com/guardian/mapper/AuditTaskMapper.java",
            "backend/src/main/java/com/guardian/mapper/AuditRecordMapper.java",
            "backend/src/main/java/com/guardian/mapper/WarningRecordMapper.java",
            "backend/src/main/java/com/guardian/vo/AuditTaskDetailVO.java",
            "backend/src/main/java/com/guardian/vo/DashboardStatsVO.java",
        ],
        "tables": ["audit_task", "audit_record", "detect_task", "detect_result", "warning_record"],
        "apis": ["/audit/tasks", "/audit/task-details", "/audit/handle", "/audit/records", "/dashboard/stats", "/warnings"],
    },
    {
        "member": "成员三",
        "role_name": "数据安全员",
        "username": "security",
        "password": "123456",
        "role_code": "SECURITY_OFFICER",
        "home": "/ai-review",
        "responsibility": "负责企业数据安全治理，维护敏感词库、执行 AI 审核、管理风险预警和风险等级规则。",
        "functions": [
            "AI 文本审核：输入审核场景、资料类型、发布范围和内容，调用 DeepSeek 或本地兜底策略给出审核建议。",
            "AI 文件审核：上传文件后抽取文本并执行 AI 审核。",
            "敏感词库管理：维护敏感词、分类、风险等级和启用状态。",
            "风险预警管理：查看和处理系统生成的高风险预警。",
            "风险等级配置：维护低、中、高、严重风险的分值区间和处理策略。",
        ],
        "frontend": [
            "frontend/src/views/AiReviewView.vue",
            "frontend/src/views/SensitiveWordView.vue",
            "frontend/src/views/WarningView.vue",
            "frontend/src/views/RiskLevelView.vue",
            "frontend/src/api/modules.js",
            "frontend/src/router/index.js",
            "frontend/src/views/LayoutView.vue",
        ],
        "backend": [
            "backend/src/main/java/com/guardian/controller/AiReviewController.java",
            "backend/src/main/java/com/guardian/controller/SensitiveWordController.java",
            "backend/src/main/java/com/guardian/controller/WarningController.java",
            "backend/src/main/java/com/guardian/controller/AdminController.java",
            "backend/src/main/java/com/guardian/service/AiReviewService.java",
            "backend/src/main/java/com/guardian/service/FileTextExtractor.java",
            "backend/src/main/java/com/guardian/service/SensitiveWordService.java",
            "backend/src/main/java/com/guardian/service/WarningService.java",
            "backend/src/main/java/com/guardian/service/AdminService.java",
            "backend/src/main/java/com/guardian/dto/AiReviewRequest.java",
            "backend/src/main/java/com/guardian/dto/SensitiveWordRequest.java",
            "backend/src/main/java/com/guardian/entity/SensitiveWord.java",
            "backend/src/main/java/com/guardian/entity/SensitiveCategory.java",
            "backend/src/main/java/com/guardian/entity/WarningRecord.java",
            "backend/src/main/java/com/guardian/entity/RiskLevelConfig.java",
            "backend/src/main/java/com/guardian/mapper/SensitiveWordMapper.java",
            "backend/src/main/java/com/guardian/mapper/SensitiveCategoryMapper.java",
            "backend/src/main/java/com/guardian/mapper/WarningRecordMapper.java",
            "backend/src/main/java/com/guardian/mapper/RiskLevelConfigMapper.java",
        ],
        "tables": ["sensitive_word", "sensitive_category", "warning_record", "risk_level_config", "detect_result"],
        "apis": ["/ai-review", "/ai-review/file", "/sensitive-words", "/sensitive-words/categories", "/warnings", "/admin/risk-levels"],
    },
    {
        "member": "成员四",
        "role_name": "系统管理员",
        "username": "admin",
        "password": "123456",
        "role_code": "ADMIN",
        "home": "/users",
        "responsibility": "负责系统运维管理，维护用户、角色、权限、日志和系统配置，保证系统基础数据和访问控制正常运行。",
        "functions": [
            "用户管理：查看系统用户，按关键字检索用户，启用或禁用账号。",
            "角色管理：维护角色编码、角色名称、说明和启用状态。",
            "权限管理：维护权限编码、权限名称、权限类型和路径，用于解释系统菜单与接口权限。",
            "日志管理：查看登录、检测、审核、风险处理等操作日志。",
            "系统配置：维护系统参数，例如阈值、开关类配置或业务提示配置。",
        ],
        "frontend": [
            "frontend/src/views/UserManageView.vue",
            "frontend/src/views/AdminConfigView.vue",
            "frontend/src/views/AdminLogView.vue",
            "frontend/src/api/modules.js",
            "frontend/src/router/index.js",
            "frontend/src/views/LayoutView.vue",
        ],
        "backend": [
            "backend/src/main/java/com/guardian/controller/UserManageController.java",
            "backend/src/main/java/com/guardian/controller/AdminController.java",
            "backend/src/main/java/com/guardian/service/UserManageService.java",
            "backend/src/main/java/com/guardian/service/AdminService.java",
            "backend/src/main/java/com/guardian/service/OperationLogService.java",
            "backend/src/main/java/com/guardian/entity/SysUser.java",
            "backend/src/main/java/com/guardian/entity/SysRole.java",
            "backend/src/main/java/com/guardian/entity/SysPermission.java",
            "backend/src/main/java/com/guardian/entity/SystemConfig.java",
            "backend/src/main/java/com/guardian/entity/OperationLog.java",
            "backend/src/main/java/com/guardian/mapper/SysUserMapper.java",
            "backend/src/main/java/com/guardian/mapper/SysRoleMapper.java",
            "backend/src/main/java/com/guardian/mapper/SysPermissionMapper.java",
            "backend/src/main/java/com/guardian/mapper/SystemConfigMapper.java",
            "backend/src/main/java/com/guardian/mapper/OperationLogMapper.java",
        ],
        "tables": ["sys_user", "sys_role", "sys_permission", "system_config", "operation_log"],
        "apis": ["/users", "/users/{id}/status", "/admin/roles", "/admin/permissions", "/admin/logs", "/admin/configs"],
    },
]

COMMON_FILES = [
    "frontend/src/main.js",
    "frontend/src/App.vue",
    "frontend/src/api/http.js",
    "frontend/src/router/index.js",
    "frontend/src/views/LoginView.vue",
    "frontend/src/views/LayoutView.vue",
    "backend/src/main/java/com/guardian/GuardianApplication.java",
    "backend/src/main/java/com/guardian/common/ApiResponse.java",
    "backend/src/main/java/com/guardian/common/BusinessException.java",
    "backend/src/main/java/com/guardian/common/GlobalExceptionHandler.java",
    "backend/src/main/java/com/guardian/config/SecurityConfig.java",
    "backend/src/main/java/com/guardian/config/CorsConfig.java",
    "backend/src/main/java/com/guardian/config/MyBatisPlusConfig.java",
    "backend/src/main/java/com/guardian/config/DataInitializer.java",
    "backend/src/main/java/com/guardian/controller/AuthController.java",
    "backend/src/main/java/com/guardian/service/AuthService.java",
    "backend/src/main/java/com/guardian/service/CaptchaService.java",
    "backend/src/main/java/com/guardian/security/JwtAuthenticationFilter.java",
    "backend/src/main/java/com/guardian/security/JwtTokenProvider.java",
    "backend/src/main/java/com/guardian/security/SecurityContext.java",
    "backend/src/main/java/com/guardian/security/LoginUserDetailsService.java",
    "backend/src/main/java/com/guardian/security/LoginUser.java",
    "database/schema.sql",
    "database/update.sql",
]


def read_text(path: str) -> str:
    file_path = ROOT / path
    if not file_path.exists():
        return "文件不存在：" + path
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return file_path.read_text(errors="replace")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, color=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "1F4E79"),
        ("Heading 3", 11, "365F91"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("智能数据守护者系统")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("四人分组用户分工与代码说明文档")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    doc.add_paragraph()


def add_code_block(doc, title: str, code: str, limit: int | None = None):
    doc.add_heading(title, level=3)
    if limit and len(code) > limit:
        code = code[:limit] + "\n\n// 由于源码较长，正文仅展示关键前段；完整代码以项目文件为准。"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(7.5)


def add_file_list_table(doc, title, files):
    doc.add_heading(title, level=3)
    rows = []
    for idx, path in enumerate(files, start=1):
        kind = "前端页面/脚本" if path.startswith("frontend") else "后端/数据库代码"
        rows.append([idx, kind, rel(path)])
    add_table(doc, ["序号", "类型", "代码文件路径"], rows, widths=[1.1, 3.0, 12.0])


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1 文档用途", level=1)
    doc.add_paragraph(
        "本文档用于说明智能数据守护者系统的四人分组开发分工。系统按照角色划分为普通员工、部门主管、数据安全员和系统管理员四类用户，"
        "每位成员负责一个用户角色的业务功能、页面展示、接口调用、后端业务处理和数据库实体维护。文档同时列出各角色的测试账号、"
        "功能描述、接口路径、前端代码文件、后端代码文件、数据库表以及关键源码摘录，方便课程设计答辩、代码讲解和后续维护。"
    )
    doc.add_paragraph(
        "说明：登录、验证码、JWT 认证、跨域配置、统一响应、异常处理、路由守卫、公共布局等属于系统公共基础代码，由四名成员共同使用；"
        "各角色章节中的代码清单只列出与该角色业务直接相关的文件。"
    )

    doc.add_heading("2 四人角色分工总表", level=1)
    rows = [
        [r["member"], r["role_name"], r["username"], r["password"], r["role_code"], r["home"], r["responsibility"]]
        for r in ROLE_DOCS
    ]
    add_table(doc, ["成员", "负责用户", "用户名", "密码", "角色编码", "登录首页", "主要职责"], rows, widths=[1.4, 2.1, 2.1, 1.8, 3.0, 2.4, 4.6])

    doc.add_heading("3 公共基础代码", level=1)
    doc.add_paragraph(
        "公共基础代码负责支撑所有用户登录、鉴权、菜单跳转、HTTP 请求封装、系统初始化和数据库建表。四个角色都依赖这些文件，"
        "但它们不归属于单一用户功能。"
    )
    add_file_list_table(doc, "3.1 公共代码文件清单", COMMON_FILES)
    add_code_block(doc, "3.2 路由权限核心代码：frontend/src/router/index.js", read_text("frontend/src/router/index.js"), limit=7000)
    add_code_block(doc, "3.3 角色菜单核心代码：frontend/src/views/LayoutView.vue", read_text("frontend/src/views/LayoutView.vue"), limit=7000)
    add_code_block(doc, "3.4 初始化四个测试用户：backend/src/main/java/com/guardian/config/DataInitializer.java", read_text("backend/src/main/java/com/guardian/config/DataInitializer.java"), limit=5000)

    doc.add_heading("4 四个用户功能与对应代码", level=1)
    for idx, role in enumerate(ROLE_DOCS, start=1):
        doc.add_heading(f"4.{idx} {role['member']}：{role['role_name']}用户", level=2)
        doc.add_paragraph(f"测试账号：{role['username']}，测试密码：{role['password']}，角色编码：{role['role_code']}，登录后默认首页：{role['home']}。")
        doc.add_paragraph(f"主要职责：{role['responsibility']}")

        doc.add_heading("功能描述", level=3)
        for item in role["functions"]:
            doc.add_paragraph(item, style="List Bullet")

        add_table(
            doc,
            ["类别", "内容"],
            [
                ["前端路由", role["home"]],
                ["接口路径", "\n".join(role["apis"])],
                ["数据库表", "\n".join(role["tables"])],
            ],
            widths=[3.0, 12.0],
        )
        add_file_list_table(doc, "对应前端代码", role["frontend"])
        add_file_list_table(doc, "对应后端代码", role["backend"])

        doc.add_heading("业务流程说明", level=3)
        if role["role_code"] == "EMPLOYEE":
            doc.add_paragraph(
                "员工从检测页面提交文本或文件后，前端调用 detectApi.detectText 或 detectApi.detectFile。后端 DetectionController 校验角色权限，"
                "DetectionService 根据正则规则和敏感词库识别手机号、身份证号、邮箱、银行卡号、地址和自定义敏感词，计算风险评分并写入 detect_task、"
                "detect_result 表。检测完成后系统自动生成 audit_task，等待部门主管审核；若风险等级较高，还会生成 warning_record。"
            )
        elif role["role_code"] == "MANAGER":
            doc.add_paragraph(
                "部门主管进入审核页面后，前端调用 auditApi.taskDetails 获取待审核任务明细。主管根据检测结果、风险评分、脱敏文本和 AI 建议判断是否通过。"
                "提交审核意见后，AuditService 更新 audit_task 和 detect_task 状态，并写入 audit_record。员工再次查看历史记录时，可以看到主管通过或驳回结果。"
            )
        elif role["role_code"] == "SECURITY_OFFICER":
            doc.add_paragraph(
                "数据安全员既可以直接提交文本进行 AI 审核，也可以上传文件进行 AI 审核。AiReviewController 对文件调用 FileTextExtractor 抽取文本，"
                "再由 AiReviewService 组装审核提示词并调用 DeepSeek 接口；当接口不可用时，系统使用本地审核建议兜底。同时，数据安全员维护敏感词库、"
                "风险预警和风险等级配置，为员工检测和主管审核提供规则基础。"
            )
        else:
            doc.add_paragraph(
                "系统管理员维护平台运行所需的基础数据，包括用户、角色、权限、系统配置和操作日志。用户管理由 UserManageController 与 UserManageService 实现；"
                "角色、权限、配置和日志由 AdminController 与 AdminService 统一处理。权限管理用于描述系统菜单、接口或操作级权限，便于说明不同角色为什么看到不同功能。"
            )

    doc.add_heading("5 分角色关键代码附录", level=1)
    for role in ROLE_DOCS:
        doc.add_heading(f"{role['member']}：{role['role_name']}关键代码", level=2)
        key_files = []
        if role["role_code"] == "EMPLOYEE":
            key_files = [
                "frontend/src/views/DetectionView.vue",
                "frontend/src/views/HistoryView.vue",
                "backend/src/main/java/com/guardian/controller/DetectionController.java",
                "backend/src/main/java/com/guardian/service/DetectionService.java",
            ]
        elif role["role_code"] == "MANAGER":
            key_files = [
                "frontend/src/views/AuditView.vue",
                "frontend/src/views/ManagerStatsView.vue",
                "backend/src/main/java/com/guardian/controller/AuditController.java",
                "backend/src/main/java/com/guardian/service/AuditService.java",
            ]
        elif role["role_code"] == "SECURITY_OFFICER":
            key_files = [
                "frontend/src/views/AiReviewView.vue",
                "frontend/src/views/SensitiveWordView.vue",
                "backend/src/main/java/com/guardian/controller/AiReviewController.java",
                "backend/src/main/java/com/guardian/service/AiReviewService.java",
                "backend/src/main/java/com/guardian/controller/SensitiveWordController.java",
            ]
        else:
            key_files = [
                "frontend/src/views/UserManageView.vue",
                "frontend/src/views/AdminConfigView.vue",
                "frontend/src/views/AdminLogView.vue",
                "backend/src/main/java/com/guardian/controller/UserManageController.java",
                "backend/src/main/java/com/guardian/controller/AdminController.java",
                "backend/src/main/java/com/guardian/service/AdminService.java",
            ]
        for path in key_files:
            add_code_block(doc, rel(path), read_text(path), limit=9000)

    doc.add_heading("6 答辩讲解建议", level=1)
    add_table(
        doc,
        ["成员", "讲解重点"],
        [
            ["成员一", "演示 employee 登录，提交文本检测和文件检测，说明检测结果如何进入员工历史记录。"],
            ["成员二", "演示 manager 登录，处理员工检测申请，说明通过或驳回后员工页面如何同步显示审核结果。"],
            ["成员三", "演示 security 登录，执行 AI 文本/文件审核，维护敏感词、预警和风险等级配置。"],
            ["成员四", "演示 admin 登录，管理用户、角色、权限、日志和系统配置，说明权限管理用于支撑角色功能隔离。"],
        ],
        widths=[2.4, 13.0],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("智能数据守护者系统四人分工与代码说明文档")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
