from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_requested_template_report as base


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
SCREENSHOT_DIR = REPORT_DIR / "assets" / "screenshots"
DIAGRAM_DIR = REPORT_DIR / "assets" / "detailed-design-diagrams"
OUT = REPORT_DIR / "智能数据守护者系统课程设计报告-详细设计无附录版.docx"


FEATURES = [
    {
        "module": "登录注册模块",
        "title": "验证码获取",
        "actor": "未登录用户",
        "desc": "验证码获取功能用于降低自动化脚本暴力登录风险。用户进入登录页后，前端调用验证码接口，后端生成验证码标识和图片内容并返回给前端。登录提交时，验证码标识与用户输入一起提交，后端先校验验证码，再校验账号与密码。",
        "steps": ["访问登录页面", "前端请求验证码接口", "后端生成验证码文本与标识", "返回验证码图片", "用户输入验证码", "登录时校验验证码"],
        "decision": ("验证码是否正确", "继续校验账号", "提示验证码错误"),
        "view": "LoginView.vue",
        "api": "/auth/captcha",
        "service": "CaptchaService",
        "tables": ["Redis/内存验证码缓存"],
        "inputs": "无",
        "outputs": "captchaKey、captchaImage",
    },
    {
        "module": "登录注册模块",
        "title": "用户登录",
        "actor": "系统用户",
        "desc": "用户登录功能是系统访问控制入口。用户输入用户名、密码和验证码后提交登录请求，后端依次校验验证码、账号是否存在、账号是否启用、密码是否匹配。校验通过后，系统生成JWT令牌并返回用户角色，前端根据角色跳转到对应首页。",
        "steps": ["填写用户名密码验证码", "提交登录请求", "校验验证码", "查询用户与账号状态", "BCrypt校验密码", "生成JWT令牌", "返回用户角色并跳转首页"],
        "decision": ("账号密码是否有效", "进入角色首页", "返回登录失败提示"),
        "view": "LoginView.vue",
        "api": "/auth/login",
        "service": "AuthService、JwtTokenProvider",
        "tables": ["sys_user", "operation_log"],
        "inputs": "username、password、captchaKey、captchaCode",
        "outputs": "token、userInfo、roleCode",
    },
    {
        "module": "登录注册模块",
        "title": "用户注册",
        "actor": "新用户",
        "desc": "用户注册功能用于创建普通业务账号。前端收集用户名、密码、真实姓名、手机号和邮箱等信息，后端校验用户名是否重复，对密码进行BCrypt加密后保存。课程设计中系统也会在首次启动时初始化四个演示账号，方便不同角色测试。",
        "steps": ["填写注册信息", "提交注册请求", "校验用户名唯一性", "加密保存密码", "写入用户表", "返回注册结果"],
        "decision": ("用户名是否已存在", "提示更换用户名", "创建用户"),
        "view": "LoginView.vue",
        "api": "/auth/register",
        "service": "AuthService",
        "tables": ["sys_user"],
        "inputs": "username、password、realName、phone、email",
        "outputs": "注册成功或失败提示",
    },
    {
        "module": "登录注册模块",
        "title": "个人资料与密码维护",
        "actor": "已登录用户",
        "desc": "个人资料与密码维护功能面向所有角色。用户进入个人中心后可以查看账号基本信息，修改真实姓名、手机号、邮箱等资料，也可以提交旧密码和新密码完成密码变更。后端通过当前JWT识别用户身份，避免用户越权修改他人资料。",
        "steps": ["进入个人中心", "加载当前用户资料", "修改基本资料或密码", "提交保存请求", "后端校验当前用户", "更新用户信息", "返回保存结果"],
        "decision": ("旧密码是否正确", "更新密码", "提示密码错误"),
        "view": "ProfileView.vue",
        "api": "/auth/profile、/auth/password",
        "service": "AuthService",
        "tables": ["sys_user", "operation_log"],
        "inputs": "profile字段、oldPassword、newPassword",
        "outputs": "更新后的个人资料、保存提示",
    },
    {
        "module": "普通员工模块",
        "title": "文本检测",
        "actor": "普通员工",
        "desc": "文本检测功能是员工提交检测内容的主要入口。员工填写任务名称和文本内容后提交，后端使用正则规则识别手机号、身份证号、邮箱、银行卡号和地址，并结合敏感词库识别企业自定义风险词。系统根据命中项计算风险评分，生成风险等级、命中摘要、脱敏文本和AI建议。",
        "steps": ["进入检测页面", "选择文本检测", "填写任务名称和文本内容", "提交检测请求", "执行规则与敏感词检测", "计算风险评分", "生成检测结果与审核任务"],
        "decision": ("是否命中高风险", "生成预警记录", "仅保存检测结果"),
        "view": "DetectionView.vue",
        "api": "/detect/text",
        "service": "DetectionService、AiReviewService",
        "tables": ["detect_task", "detect_result", "audit_task", "warning_record", "sensitive_word"],
        "inputs": "taskName、content",
        "outputs": "riskLevel、riskScore、hitSummary、desensitizedContent、aiSuggestion",
    },
    {
        "module": "普通员工模块",
        "title": "文件检测",
        "actor": "普通员工",
        "desc": "文件检测功能用于处理员工上传的txt、pdf、doc和docx文件。前端通过上传组件构造表单数据，后端先由文件文本抽取服务读取文件内容，再进入与文本检测一致的风险识别流程。该功能解决了实际工作中文档类资料需要发布前检查的问题。",
        "steps": ["选择文件检测", "上传待检测文件", "后端接收MultipartFile", "抽取文件文本内容", "执行敏感信息检测", "生成结果和审核任务", "页面展示检测结论"],
        "decision": ("文件是否可解析", "继续检测", "提示重新上传"),
        "view": "DetectionView.vue",
        "api": "/detect/file",
        "service": "FileTextExtractor、DetectionService",
        "tables": ["detect_task", "detect_result", "audit_task", "warning_record"],
        "inputs": "taskName、file",
        "outputs": "文件检测结果、脱敏内容、AI建议",
    },
    {
        "module": "普通员工模块",
        "title": "检测结果查看",
        "actor": "普通员工",
        "desc": "检测结果查看功能用于让员工在提交后立即理解风险原因。系统返回风险等级、风险评分、命中摘要、脱敏内容和AI建议，员工可以根据这些信息决定是否调整内容。该功能强调结果解释性，不只告诉用户有风险，还说明风险来自哪些字段。",
        "steps": ["提交检测内容", "后端返回检测结果", "前端渲染风险等级", "展示命中摘要", "展示脱敏结果", "展示AI建议"],
        "decision": ("是否需要修改内容", "员工重新编辑提交", "进入主管审核流程"),
        "view": "DetectionView.vue",
        "api": "/detect/text、/detect/file",
        "service": "DetectionService",
        "tables": ["detect_result"],
        "inputs": "检测响应数据",
        "outputs": "可视化检测结论",
    },
    {
        "module": "普通员工模块",
        "title": "历史记录查看",
        "actor": "普通员工",
        "desc": "历史记录查看功能用于查询当前员工提交过的全部检测任务。后端根据JWT中的用户编号筛选detect_task，并关联detect_result和audit_task，返回任务名称、内容类型、风险等级、风险评分、审核状态和审核意见。员工可通过该页面确认主管是否通过或驳回。",
        "steps": ["进入历史记录页面", "请求个人历史接口", "后端按当前用户查询任务", "关联检测结果与审核任务", "返回历史记录列表", "员工查看审核状态"],
        "decision": ("审核是否完成", "显示审核意见", "显示待审核"),
        "view": "HistoryView.vue",
        "api": "/detect/history",
        "service": "DetectionService",
        "tables": ["detect_task", "detect_result", "audit_task"],
        "inputs": "当前登录用户ID",
        "outputs": "检测历史列表",
    },
    {
        "module": "数据安全员模块",
        "title": "AI文本审核",
        "actor": "数据安全员",
        "desc": "AI文本审核功能用于对企业待发布文本进行安全复核。安全员选择审核场景、资料类型和发布范围，输入文本后提交，后端组装提示词调用DeepSeek接口，要求模型输出风险等级、发布建议、风险依据、修改建议和审核结论。审核结果会保存到AI审核记录表。",
        "steps": ["进入AI审核页面", "选择审核场景", "填写资料类型与发布范围", "输入待审核文本", "调用AI审核服务", "解析结构化审核结论", "保存AI审核记录"],
        "decision": ("AI接口是否可用", "返回DeepSeek结论", "使用本地兜底建议"),
        "view": "AiReviewView.vue",
        "api": "/ai-review",
        "service": "AiReviewService",
        "tables": ["ai_review_record", "operation_log"],
        "inputs": "reviewScene、dataType、publishScope、content",
        "outputs": "结构化AI审核结论",
    },
    {
        "module": "数据安全员模块",
        "title": "AI文件审核",
        "actor": "数据安全员",
        "desc": "AI文件审核功能面向合同、技术文档、财务资料等文件类内容。安全员上传文件后，后端抽取文本并复用AI审核服务，审核结论保存时会同时记录文件名和内容类型。该功能与员工文件检测不同，重点不是规则命中评分，而是让安全员获得面向发布决策的综合审核意见。",
        "steps": ["选择AI文件审核", "上传待审核文件", "抽取文件文本", "补充审核上下文", "调用AI审核服务", "保存文件审核记录", "显示审核结论"],
        "decision": ("文件内容是否抽取成功", "进入AI审核", "提示文件无法解析"),
        "view": "AiReviewView.vue",
        "api": "/ai-review/file",
        "service": "FileTextExtractor、AiReviewService",
        "tables": ["ai_review_record"],
        "inputs": "reviewScene、dataType、publishScope、file",
        "outputs": "文件AI审核结论",
    },
    {
        "module": "数据安全员模块",
        "title": "AI审核记录管理",
        "actor": "数据安全员",
        "desc": "AI审核记录管理功能用于复盘历史AI审核。每次文本或文件AI审核都会保存为记录，安全员可以查看审核场景、资料类型、发布范围、内容类型、文件名、原始内容、审核结论和时间。该功能使AI审核从一次性页面结果变为可追踪的安全审计数据。",
        "steps": ["进入AI审核记录页面", "请求审核记录接口", "后端查询记录列表", "前端展示摘要字段", "点击详情", "查看原文和审核结论"],
        "decision": ("是否需要复盘详情", "打开详情弹窗", "继续浏览列表"),
        "view": "AiReviewRecordsView.vue",
        "api": "/ai-review/records",
        "service": "AiReviewService",
        "tables": ["ai_review_record"],
        "inputs": "当前安全员身份",
        "outputs": "AI审核记录列表与详情",
    },
    {
        "module": "数据安全员模块",
        "title": "敏感词库管理",
        "actor": "数据安全员",
        "desc": "敏感词库管理功能用于维护企业个性化安全规则。安全员可以按关键字查询敏感词，新增或修改敏感词所属分类、风险等级和启用状态，也可以删除过期词条。检测服务会读取启用状态的敏感词，因此词库变化会直接影响后续检测结果。",
        "steps": ["进入敏感词库页面", "加载分类和词条", "新增或编辑敏感词", "校验分类是否存在", "保存敏感词配置", "后续检测使用新规则"],
        "decision": ("敏感词分类是否有效", "保存词条", "提示分类错误"),
        "view": "SensitiveWordView.vue",
        "api": "/sensitive-words",
        "service": "SensitiveWordService",
        "tables": ["sensitive_category", "sensitive_word", "operation_log"],
        "inputs": "categoryId、word、riskLevel、enabled",
        "outputs": "敏感词列表、保存结果",
    },
    {
        "module": "数据安全员模块",
        "title": "风险预警管理",
        "actor": "数据安全员",
        "desc": "风险预警管理功能用于处理高风险或严重风险检测任务。系统在员工检测命中高风险时自动生成预警记录，安全员进入预警页面后可以查看预警等级、内容、关联任务和处理状态，并对未处理预警进行处理，形成风险闭环。",
        "steps": ["进入风险预警页面", "加载预警列表", "查看预警等级和内容", "选择未处理预警", "执行处理操作", "更新预警状态"],
        "decision": ("预警是否已处理", "展示已处理状态", "允许执行处理"),
        "view": "WarningView.vue",
        "api": "/warnings、/warnings/{id}/handle",
        "service": "WarningService",
        "tables": ["warning_record", "operation_log"],
        "inputs": "warningId",
        "outputs": "预警列表、处理结果",
    },
    {
        "module": "数据安全员模块",
        "title": "风险等级配置",
        "actor": "数据安全员",
        "desc": "风险等级配置功能用于维护风险分值区间和处理规则。系统通过风险评分映射低风险、中风险、高风险和严重风险，安全员可以维护等级名称、最低分、最高分、处理规则和启用状态。该功能保证风险判定规则可以随着企业制度变化而调整。",
        "steps": ["进入风险等级配置", "加载等级列表", "编辑分值区间和处理规则", "提交保存", "后端写入配置表", "后续检测参考规则说明"],
        "decision": ("分值区间是否合理", "保存配置", "提示重新填写"),
        "view": "RiskLevelView.vue",
        "api": "/admin/risk-levels",
        "service": "AdminService",
        "tables": ["risk_level_config"],
        "inputs": "levelCode、levelName、minScore、maxScore、handleRule",
        "outputs": "风险等级配置列表",
    },
    {
        "module": "系统管理员模块",
        "title": "用户管理",
        "actor": "系统管理员",
        "desc": "用户管理功能用于维护系统账号。管理员可以查看用户列表，按用户名或真实姓名检索用户，并启用或禁用账号。账号禁用后用户无法继续登录，从而实现运维层面的访问控制。",
        "steps": ["进入用户管理页面", "查询用户列表", "输入关键字筛选", "选择启用或禁用", "后端更新用户状态", "刷新列表"],
        "decision": ("是否确认变更状态", "更新账号状态", "取消操作"),
        "view": "UserManageView.vue",
        "api": "/users、/users/{id}/status",
        "service": "UserManageService",
        "tables": ["sys_user", "operation_log"],
        "inputs": "keyword、userId、status",
        "outputs": "用户列表、状态更新结果",
    },
    {
        "module": "系统管理员模块",
        "title": "角色管理",
        "actor": "系统管理员",
        "desc": "角色管理功能用于维护系统角色基础数据。管理员可以查看角色列表，保存角色编码、角色名称和说明。系统当前包含普通员工、部门主管、数据安全员和系统管理员四类角色，角色数据为菜单隔离和接口权限判断提供业务依据。",
        "steps": ["进入角色管理页面", "加载角色列表", "编辑角色信息", "提交角色配置", "后端保存角色", "刷新角色列表"],
        "decision": ("角色编码是否重复", "提示修改编码", "保存角色"),
        "view": "AdminConfigView.vue",
        "api": "/admin/roles",
        "service": "AdminService",
        "tables": ["sys_role"],
        "inputs": "roleCode、roleName、description",
        "outputs": "角色列表、保存提示",
    },
    {
        "module": "系统管理员模块",
        "title": "权限管理",
        "actor": "系统管理员",
        "desc": "权限管理功能用于维护系统菜单或接口权限说明。管理员可以维护权限编码、权限名称和菜单路径，例如文本检测、审核申请、敏感词库管理、用户管理等权限。该功能在课程设计中用于展示RBAC权限模型，使不同角色功能隔离具备数据依据。",
        "steps": ["进入权限管理页面", "加载权限列表", "编辑权限编码和路径", "提交权限配置", "后端保存权限", "刷新权限列表"],
        "decision": ("权限编码是否重复", "提示修改编码", "保存权限"),
        "view": "AdminConfigView.vue",
        "api": "/admin/permissions",
        "service": "AdminService",
        "tables": ["sys_permission"],
        "inputs": "permissionCode、permissionName、menuPath",
        "outputs": "权限列表、保存提示",
    },
    {
        "module": "系统管理员模块",
        "title": "日志管理",
        "actor": "系统管理员",
        "desc": "日志管理功能用于查看系统关键操作。登录、检测、审核、敏感词维护、风险预警处理和配置维护等行为会形成操作日志。管理员通过日志页面可以查看用户编号、日志类型、操作名称、详情、IP地址和创建时间，为问题追踪和安全审计提供依据。",
        "steps": ["进入日志管理页面", "请求日志接口", "后端按时间倒序查询日志", "前端展示日志列表", "管理员查看操作详情"],
        "decision": ("是否发现异常操作", "定位用户和时间", "继续监控"),
        "view": "AdminLogView.vue",
        "api": "/admin/logs",
        "service": "AdminService、OperationLogService",
        "tables": ["operation_log"],
        "inputs": "无",
        "outputs": "系统操作日志列表",
    },
    {
        "module": "系统管理员模块",
        "title": "系统配置",
        "actor": "系统管理员",
        "desc": "系统配置功能用于维护平台运行参数，例如JWT有效期、DeepSeek模型名称、文件上传大小等。管理员通过配置页面调整参数，后端保存到system_config表。配置数据使系统具备一定的可维护性，避免所有参数都硬编码在业务代码中。",
        "steps": ["进入系统配置页面", "加载配置列表", "编辑配置值和说明", "提交配置", "后端保存配置", "页面刷新展示新值"],
        "decision": ("配置值是否为空", "保存配置", "提示补全配置"),
        "view": "AdminConfigView.vue",
        "api": "/admin/configs",
        "service": "AdminService",
        "tables": ["system_config"],
        "inputs": "configKey、configValue、description",
        "outputs": "系统配置列表、保存提示",
    },
]


def clean_name(text: str) -> str:
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff]+", "_", text)


def title(draw: ImageDraw.ImageDraw, text: str) -> None:
    draw.text((48, 28), text, font=base.font(30, True), fill="#0f172a")
    draw.line((48, 78, 1152, 78), fill="#cbd5e1", width=2)


def uml_activity(feature: dict, idx: int) -> Path:
    path = DIAGRAM_DIR / f"{idx:02d}_{clean_name(feature['title'])}_activity.png"
    img = Image.new("RGB", (1200, 1050), "#f8fafc")
    d = ImageDraw.Draw(img)
    title(d, f"UML活动图：{feature['title']}")
    d.text((65, 95), f"参与者：{feature['actor']}", font=base.font(18), fill="#334155")
    x, y = 420, 150
    d.ellipse((570, y, 630, y + 60), fill="#0f172a")
    prev = (600, y + 60)
    y += 95
    steps = feature["steps"]
    decision_after = max(2, min(4, len(steps) - 2))
    for step_idx, step in enumerate(steps):
        base.node(d, (x, y, x + 360, y + 66), step, "#ffffff", "#2563eb", 18)
        base.arrow(d, prev, (600, y), "#334155", 3)
        prev = (600, y + 66)
        y += 98
        if step_idx == decision_after:
            question, yes, no = feature["decision"]
            base.diamond(d, 600, y + 48, 300, 112, question)
            base.arrow(d, prev, (600, y - 8), "#334155", 3)
            base.node(d, (790, y + 12, 1110, y + 84), yes, "#dcfce7", "#16a34a", 16)
            base.node(d, (90, y + 12, 410, y + 84), no, "#fee2e2", "#dc2626", 16)
            base.arrow(d, (750, y + 48), (790, y + 48), "#16a34a", 3)
            base.arrow(d, (450, y + 48), (410, y + 48), "#dc2626", 3)
            d.text((752, y + 22), "是", font=base.font(16, True), fill="#166534")
            d.text((420, y + 22), "否", font=base.font(16, True), fill="#991b1b")
            prev = (600, y + 104)
            y += 145
    d.ellipse((568, y, 632, y + 64), outline="#0f172a", width=4)
    d.ellipse((581, y + 13, 619, y + 51), fill="#0f172a")
    base.arrow(d, prev, (600, y), "#334155", 3)
    img.save(path)
    return path


def business_model(feature: dict, idx: int) -> Path:
    path = DIAGRAM_DIR / f"{idx:02d}_{clean_name(feature['title'])}_business.png"
    img = Image.new("RGB", (1400, 760), "#f8fafc")
    d = ImageDraw.Draw(img)
    title(d, f"业务模型图：{feature['title']}")
    boxes = [
        ("业务角色", feature["actor"], 55, 190, "#eff6ff", "#2563eb"),
        ("前端页面", feature["view"], 280, 190, "#ffffff", "#64748b"),
        ("接口地址", feature["api"], 505, 190, "#ffffff", "#64748b"),
        ("业务服务", feature["service"], 730, 190, "#ffffff", "#64748b"),
        ("数据对象", "\n".join(feature["tables"]), 955, 190, "#fefce8", "#ca8a04"),
    ]
    centers = []
    for label, value, x, y, fill, outline in boxes:
        d.rounded_rectangle((x, y, x + 190, y + 220), radius=18, fill=fill, outline=outline, width=3)
        d.rectangle((x, y, x + 190, y + 42), fill=outline)
        base.centered(d, (x, y, x + 190, y + 42), label, base.font(18, True), "#ffffff")
        base.centered(d, (x + 8, y + 54, x + 182, y + 210), value, base.font(17), "#0f172a")
        centers.append((x + 190, y + 110))
    for i in range(len(centers) - 1):
        base.arrow(d, centers[i], (boxes[i + 1][2], boxes[i + 1][3] + 110), "#334155", 3)
    d.text((70, 480), "输入数据", font=base.font(22, True), fill="#0f172a")
    base.node(d, (180, 455, 610, 545), feature["inputs"], "#ffffff", "#2563eb", 16)
    d.text((720, 480), "输出结果", font=base.font(22, True), fill="#0f172a")
    base.node(d, (830, 455, 1310, 545), feature["outputs"], "#ffffff", "#16a34a", 16)
    base.arrow(d, (610, 500), (830, 500), "#64748b", 2)
    d.text((70, 610), "模型说明：业务角色通过前端页面触发接口请求，后端控制器进入业务服务层，服务层完成校验、计算、保存或查询，并最终读写对应数据对象。", font=base.font(18), fill="#334155")
    img.save(path)
    return path


def generate_detailed_diagrams() -> dict[str, tuple[Path, Path]]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    result = {}
    for idx, feature in enumerate(FEATURES, start=1):
        result[feature["title"]] = (uml_activity(feature, idx), business_model(feature, idx))
    return result


def add_heading_numbered(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_function_detail(doc: Document, feature: dict, diagrams: dict[str, tuple[Path, Path]], fig_no: list[int]) -> None:
    doc.add_heading(feature["title"], level=3)
    base.add_para(doc, feature["desc"])
    for paragraph in detailed_paragraphs(feature):
        base.add_para(doc, paragraph)
    rows = [
        ["业务参与者", feature["actor"]],
        ["前端页面", feature["view"]],
        ["后端接口", feature["api"]],
        ["业务服务", feature["service"]],
        ["涉及数据", "、".join(feature["tables"])],
        ["输入数据", feature["inputs"]],
        ["输出结果", feature["outputs"]],
    ]
    base.add_table(doc, ["设计项", "详细说明"], rows, widths=[3.0, 12.5])
    activity, model = diagrams[feature["title"]]
    base.add_picture(doc, activity, f"图2-{fig_no[0]} {feature['title']}UML活动图", 5.7)
    fig_no[0] += 1
    base.add_picture(doc, model, f"图2-{fig_no[0]} {feature['title']}业务模型图", 5.9)
    fig_no[0] += 1


def detailed_paragraphs(feature: dict) -> list[str]:
    title = feature["title"]
    module = feature["module"]
    actor = feature["actor"]
    api = feature["api"]
    service = feature["service"]
    tables = "、".join(feature["tables"])
    inputs = feature["inputs"]
    outputs = feature["outputs"]
    view = feature["view"]
    steps = "、".join(feature["steps"])

    common = [
        f"从功能目标看，{title}不是孤立页面操作，而是{module}中的一个完整业务动作。该动作由{actor}发起，前端页面为{view}，后端接口为{api}，核心处理服务为{service}。设计时需要保证页面输入、接口参数、业务服务和数据表字段之间保持一致，使用户在页面看到的处理结果能够追溯到后端记录。",
        f"从处理流程看，{title}的主要步骤包括{steps}。这些步骤在UML活动图中用初始节点、动作节点、判断节点和终止节点表示，能够体现用户操作、系统处理、条件判断和结果返回之间的先后关系。该流程设计的重点是让正常路径清晰，同时把失败路径和提示路径单独画出，避免业务只描述成功情况。",
        f"从数据流转看，该功能的输入数据包括{inputs}，输出结果包括{outputs}，涉及的数据对象包括{tables}。前端负责收集输入并展示输出，后端负责参数校验、权限校验、业务计算和数据持久化。对于需要保存记录的功能，系统不仅返回页面结果，还会写入数据库，保证后续查询、审计或复盘时有据可查。",
        f"从权限控制看，{title}必须限制在{actor}的职责范围内。前端通过路由和菜单控制用户可见功能，后端通过接口权限注解和JWT中的角色信息进行再次校验。这样即使用户手动输入地址或直接请求接口，也不能越权访问其他角色的业务数据。",
        f"从异常处理看，{title}需要考虑输入为空、参数格式错误、数据不存在、文件无法解析、接口调用失败或数据库保存失败等情况。页面应给出明确提示，后端应使用统一响应结构返回错误信息。对于影响安全审计的功能，还需要写入操作日志或保存业务记录，避免异常发生后无法定位问题。",
        f"从测试关注点看，{title}需要同时验证正常路径、失败路径和权限边界。正常路径用于确认功能能够完成业务目标，失败路径用于确认系统能够给出合理提示，权限边界用于确认其他角色不能绕过菜单直接调用该功能。该测试思路与活动图中的判断节点保持一致。",
        f"从模块协作看，{title}与系统其他功能存在数据或流程衔接关系。前一功能产生的数据可能成为本功能输入，本功能保存的结果也可能被历史记录、日志审计、预警处理或后台管理页面继续使用。因此详细设计中需要同时描述页面行为、接口行为和数据行为，不能只停留在界面说明。",
    ]

    specific = {
        "验证码获取": [
            "验证码获取的关键规则是验证码必须与一次登录会话绑定，并设置有效期。验证码图片返回给前端后，用户输入的验证码不能只在前端判断，而必须交给后端校验。这样可以避免绕过前端页面直接调用登录接口，提高登录入口的基础防护能力。",
            "业务模型中验证码缓存不属于永久业务表，但它是登录流程中的临时安全对象。验证码生成、展示、提交和校验构成短生命周期数据流，校验成功或过期后应及时失效，避免同一验证码被重复使用。",
        ],
        "用户登录": [
            "用户登录的核心业务规则是先校验验证码，再校验账号状态，最后校验密码。这样的顺序能够减少无效账号密码校验次数，也能在账号被禁用时直接阻断登录。密码校验必须使用BCrypt等安全算法，不能明文比较或明文存储。",
            "登录成功后返回JWT令牌和角色编码，前端根据角色编码跳转到对应首页。普通员工进入检测页面，数据安全员进入AI审核页面，系统管理员进入用户管理页面。登录失败时不应暴露过多账号细节，只给出统一提示，避免攻击者枚举账号状态。",
        ],
        "用户注册": [
            "用户注册需要重点处理用户名唯一性和密码安全。后端保存前必须查询用户名是否已存在，避免重复账号造成登录混乱。密码必须加密保存，注册接口不应返回密码字段，也不应在日志中输出原始密码。",
            "在课程设计演示中，系统初始化了四个固定角色账号；注册功能则用于展示用户扩展能力。真实环境中注册通常需要管理员审核或默认赋予普通员工角色，避免用户自行注册成为高权限角色。",
        ],
        "个人资料与密码维护": [
            "个人资料维护强调当前用户只能修改自己的资料。后端不应信任前端传入的用户编号，而应从JWT安全上下文中获取当前用户ID。密码修改必须校验旧密码，旧密码错误时不能更新新密码。",
            "该功能虽然不直接参与检测审核流程，但它关系到账号安全和用户信息准确性。手机号、邮箱等资料可以用于后续通知和审计，因此保存时需要进行基本格式校验，避免无效数据进入用户表。",
        ],
        "文本检测": [
            "文本检测的核心业务规则是风险识别和风险评分。系统通过正则表达式识别格式化敏感信息，通过敏感词库识别业务敏感词，再按照命中项风险等级累加分值，得到最终风险等级。风险分值上限应受控，避免极端文本导致评分失真。",
            "文本检测完成后必须自动创建审核任务，这体现了检测与审核的流程衔接。员工端只负责提交和查看结果，是否允许后续发布需要主管或安全治理流程判断。高风险内容还要生成预警记录，使安全员能够及时处理。",
        ],
        "文件检测": [
            "文件检测的关键难点是文件内容抽取。不同文件格式的文本结构不同，后端需要根据扩展名或文件类型选择解析方式。解析失败时不能继续生成错误检测结果，而应返回明确提示，让员工重新上传可解析文件。",
            "文件检测与文本检测复用同一套风险识别逻辑，这样可以保证规则一致。无论内容来自直接输入还是文件抽取，最终都应生成检测任务、检测结果和审核任务，便于历史记录统一展示。",
        ],
        "检测结果查看": [
            "检测结果查看强调结果解释性。系统不能只显示风险等级，还要展示命中摘要和脱敏内容，使员工知道风险来自哪些字段。AI建议用于辅助员工修改内容，但不替代主管审核结论。",
            "该功能的页面展示应区分原始内容、命中信息和脱敏结果。对于银行卡号、手机号、身份证号等字段，脱敏展示可以降低二次泄露风险。员工根据提示修改后，可以重新提交检测任务。",
        ],
        "历史记录查看": [
            "历史记录查看必须按当前用户过滤，员工只能查看自己提交的检测任务。后端需要关联检测结果和审核任务，否则员工只能看到任务本身，无法知道风险结果和主管审核意见。",
            "该功能是员工理解流程状态的窗口。待审核、已通过、已驳回等状态需要清晰展示；驳回时应显示审核意见，帮助员工知道需要修改的方向。历史记录也能证明系统流程是闭环而非一次性检测。",
        ],
        "AI文本审核": [
            "AI文本审核的核心是提示词构造。系统把审核场景、资料类型、发布范围和待审核内容组合成上下文，要求模型输出风险等级、发布建议、风险依据、修改建议和审核结论。这样可以减少模型输出随意性，提高审核结果可读性。",
            "AI文本审核结果需要保存到AI审核记录表。保存内容包括安全员编号、审核场景、资料类型、发布范围、内容类型、原始文本和审核结论。这样安全员后续可以复盘历史审核，也便于管理员或老师检查流程是否真实发生。",
        ],
        "AI文件审核": [
            "AI文件审核先进行文件文本抽取，再进入AI审核流程。它与员工文件检测的区别在于：员工文件检测偏向规则识别和风险评分，安全员AI文件审核偏向发布决策和综合建议。两者可以互补，而不是重复功能。",
            "保存文件审核记录时需要记录文件名和内容类型，方便后续追溯审核来源。如果文件解析失败，系统不应调用AI接口，否则模型得到的内容为空或不完整，会导致审核结论不可靠。",
        ],
        "AI审核记录管理": [
            "AI审核记录管理是新增的数据安全员功能，用于解决AI审核结果只在页面临时显示的问题。通过记录列表和详情弹窗，安全员可以查看历史审核内容、文件名、审核场景、发布范围和审核结论。",
            "该功能提升了系统的审计能力。对于课程答辩，AI审核记录可以证明文本审核和文件审核确实落库；对于真实业务，记录可以作为后续复盘和责任追踪依据。",
        ],
        "敏感词库管理": [
            "敏感词库管理使安全规则具备可维护性。企业敏感信息不一定都有固定格式，例如客户名单、报价策略、源代码、接口密钥等内容需要通过业务词库补充识别。安全员维护词库后，后续员工检测会自动使用新规则。",
            "保存敏感词时必须校验分类是否存在，避免产生无分类的孤立词条。风险等级字段也必须与系统风险等级保持一致，否则检测服务无法正确计算评分或展示风险。",
        ],
        "风险预警管理": [
            "风险预警管理连接检测结果和安全治理。检测服务发现高风险或严重风险后自动生成预警，安全员在预警页面查看并处理。处理状态更新后，系统能够区分未处理风险和已闭环风险。",
            "预警处理应写入日志，因为它属于安全管理动作。日志中记录预警编号、处理人、处理时间和处理动作，能够帮助管理员追踪风险处理责任。",
        ],
        "风险等级配置": [
            "风险等级配置用于解释分数与处理策略之间的关系。低风险可以记录后继续流转，中风险建议脱敏后审核，高风险需要生成预警，严重风险应阻断发布并要求复核。配置页面使这些规则可以被展示和维护。",
            "分值区间需要避免交叉和空缺，否则同一风险评分可能匹配多个等级或无法匹配等级。保存时应校验最低分、最高分和启用状态，保证规则可用。",
        ],
        "用户管理": [
            "用户管理是系统管理员最基础的运维功能。管理员可以禁用异常账号，禁用后该账号即使知道密码也不能继续登录。用户列表中的角色、部门和状态字段可以帮助管理员判断账号归属。",
            "用户状态变更应记录操作日志，避免管理员误操作后无法追踪。对于真实系统，还可以增加新增用户、重置密码、分配角色等扩展功能。",
        ],
        "角色管理": [
            "角色管理定义系统中的职责边界。普通员工、部门主管、数据安全员和系统管理员分别对应不同菜单和接口。角色编码应保持稳定，因为前端路由和后端权限判断都依赖角色编码。",
            "角色信息修改应谨慎，尤其是角色编码。若编码变化但路由和接口权限未同步，会导致用户无法访问功能或出现权限错乱。因此角色管理更适合维护名称和说明，编码应作为核心标识。",
        ],
        "权限管理": [
            "权限管理用于说明系统中有哪些菜单和操作能力。虽然当前项目主要通过角色直接控制菜单和接口，但权限表为后续扩展细粒度RBAC提供基础。权限编码可以对应接口动作，菜单路径可以对应前端路由。",
            "权限管理的业务意义是让系统管理员能够解释为什么不同角色功能不同。课程设计中展示权限表，可以证明系统不是把页面简单隐藏，而是具备角色与权限模型基础。",
        ],
        "日志管理": [
            "日志管理面向审计和问题定位。系统运行中，登录、检测、审核、敏感词维护、预警处理、配置保存等操作都可能影响业务安全，因此需要记录用户、类型、动作、详情、IP和时间。",
            "管理员查看日志时，可以根据时间和操作类型分析异常行为。例如某用户频繁提交高风险检测，或某管理员频繁修改配置，都可以通过日志发现线索。",
        ],
        "系统配置": [
            "系统配置让部分运行参数从代码中抽离出来。JWT有效期、AI模型名称、上传文件大小等参数如果全部硬编码，后续调整需要重新发布代码；保存到配置表后，系统维护更加灵活。",
            "配置保存需要注意参数有效性。比如JWT有效期应为正数，文件大小限制不能过大，模型名称不能为空。对于关键配置，真实系统还应记录修改前后值，方便回滚。",
        ],
    }
    return common + specific.get(title, [])


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    lines = [
        "1 可行性研究与需求分析\t1",
        "1.1 选题意义\t1",
        "1.2 可行性研究\t2",
        "1.2.1 技术可行性\t2",
        "1.2.2 运行可行性\t2",
        "1.2.3 操作可行性\t2",
        "1.2.4 法律可行性\t2",
        "1.3 需求分析\t3",
        "1.3.1 安全性需求\t3",
        "1.3.2 稳定性和可维护性需求\t3",
        "1.3.3 主要功能需求\t3",
        "2 系统设计\t4",
        "2.1 设计原则\t4",
        "2.2 主要模块功能（概要设计）\t5",
        "2.3 各个模块详细设计\t7",
        "2.3.1 登录注册模块\t7",
        "2.3.2 普通员工模块\t12",
        "2.3.3 数据安全员模块\t17",
        "2.3.4 系统管理员模块\t25",
        "2.4 数据库概念结构模型\t31",
        "2.4.1 表结构设计\t31",
        "2.4.2 E-R模型\t34",
        "3 系统实现与测试\t35",
        "4 总结\t44",
    ]
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def add_front_sections(doc: Document, overview_diagrams: dict[str, Path]) -> None:
    doc.add_heading("1 可行性研究与需求分析", level=1)
    doc.add_heading("1.1 选题意义", level=2)
    for text in [
        "智能数据守护者系统以企业数据流转安全为背景，围绕员工提交内容、系统自动检测、主管审核、安全员治理和管理员运维这一完整链路展开设计。现实工作中，客户名单、合同报价、财务数据、身份证号、手机号、银行卡号、技术接口密钥等信息经常混杂在普通文档和文本中，如果缺少发布前检测与审批机制，极易造成隐私泄露、商业秘密泄露和合规风险。",
        "本系统的选题意义在于将数据安全控制前移到内容提交阶段。普通员工不需要完全依赖经验判断文本或文件是否安全，系统能够自动检测敏感内容并给出脱敏建议；数据安全员可以维护敏感词库和风险等级规则，并使用AI审核辅助判断发布风险；管理员可以维护用户、角色、权限和日志，保证系统具备可管理性。系统既服务于课程设计，也体现了企业数据安全治理中“发现风险、解释风险、审核风险、记录风险”的基本思想。",
        "从课程训练角度看，本系统覆盖Spring Boot、MyBatis Plus、MySQL、Vue3、Element Plus、JWT、文件上传、DeepSeek API、UML活动图、业务模型、E-R模型和系统测试等综合内容。通过该题目可以把需求分析、概要设计、详细设计、数据库设计、编码实现和运行测试串联起来，较完整地展示专业领域实践课程设计能力。",
    ]:
        base.add_para(doc, text)
    doc.add_heading("1.2 可行性研究", level=2)
    feasibility = [
        ("1.2.1 技术可行性", "系统采用Spring Boot 3作为后端框架，使用MyBatis Plus完成数据库访问，使用MySQL保存业务数据，前端采用Vue3、Vite和Element Plus实现页面交互。敏感信息识别采用正则规则、敏感词库和AI辅助审核结合的方式，DeepSeek接口异常时保留本地兜底建议。因此技术栈成熟、资料充足、实现路径清晰，适合课程设计周期内完成。"),
        ("1.2.2 运行可行性", "系统可在普通开发机上运行，后端通过Maven启动，前端通过Vite启动，数据库使用本地MySQL即可完成演示。系统核心业务主要是文本解析、数据库读写和少量文件内容抽取，对硬件资源要求不高。若部署到真实环境，也可以将前端、后端、数据库和AI服务配置拆分部署。"),
        ("1.2.3 操作可行性", "系统按角色组织菜单，员工、安全员、管理员进入系统后看到的功能不同。页面使用表单、表格、上传控件和弹窗组织操作，用户只需按照业务流程填写和提交即可。检测结果以风险等级、风险评分、命中摘要、脱敏内容和建议方式展示，便于非技术人员理解。"),
        ("1.2.4 法律可行性", "系统目标是降低敏感数据泄露风险，符合访问控制、最小必要、操作留痕和人工复核等安全管理原则。用户密码采用加密存储，接口按角色鉴权，敏感内容提供脱敏结果，AI审核只作为辅助建议，最终处理仍由人工确认。课程设计环境中不连接真实企业数据，具备法律可行性。"),
    ]
    for heading, text in feasibility:
        doc.add_heading(heading, level=3)
        base.add_para(doc, text)
    doc.add_heading("1.3 需求分析", level=2)
    reqs = [
        ("1.3.1 安全性需求", ["登录必须校验验证码、账号状态和密码。", "系统接口必须基于JWT识别当前用户并进行角色限制。", "敏感信息需要脱敏展示，避免检测页面再次扩大泄露范围。", "关键操作需要写入日志，便于管理员审计。", "高风险检测结果需要生成预警并进入处理流程。"]),
        ("1.3.2 稳定性和可维护性需求", ["后端代码需要按Controller、Service、Mapper、Entity、DTO、VO分层。", "AI接口不可用时应返回本地兜底建议，保证流程不中断。", "敏感词、风险等级、系统配置等规则应支持后台维护。", "数据库表保留创建时间、更新时间和逻辑删除字段。", "前端接口封装应统一处理令牌、响应和错误提示。"]),
        ("1.3.3 主要功能需求", ["登录注册模块：验证码、登录、注册、个人资料与密码维护。", "普通员工模块：文本检测、文件检测、检测结果查看、历史记录查看。", "数据安全员模块：AI文本审核、AI文件审核、AI审核记录、敏感词库、风险预警、风险等级配置。", "系统管理员模块：用户管理、角色管理、权限管理、日志管理、系统配置。", "系统测试需要覆盖各角色登录、菜单隔离、检测提交、AI审核、记录查询和后台管理。"]),
    ]
    for heading, items in reqs:
        doc.add_heading(heading, level=3)
        base.add_bullets(doc, items)

    doc.add_heading("2 系统设计", level=1)
    doc.add_heading("2.1 设计原则", level=2)
    base.add_para(doc, "系统设计遵循角色分离、流程闭环、规则可配置、接口统一和可追溯五项原则。角色分离要求不同用户只能访问职责范围内的功能；流程闭环要求检测、审核、预警、记录和日志形成完整链路；规则可配置要求敏感词库、风险等级和系统参数可以后台维护；接口统一要求前后端采用统一响应结构；可追溯要求AI审核记录、操作日志和审核记录都能长期保存。")
    base.add_picture(doc, overview_diagrams["system_data_model"], "图2-1 系统数据模型图", 6.2)
    doc.add_heading("2.2 主要模块功能（概要设计）", level=2)
    base.add_para(doc, "系统主要模块包括登录注册模块、普通员工模块、部门主管模块、数据安全员模块、系统管理员模块和公共支撑模块。虽然本次详细设计重点描述登录注册、普通员工、数据安全员和系统管理员，但概要设计中仍保留部门主管审核流程，因为员工提交的检测任务需要经过主管审核后才能形成完整业务闭环。")
    base.add_picture(doc, overview_diagrams["module_structure"], "图2-2 模块结构模型图", 6.2)
    base.add_picture(doc, overview_diagrams["module_logic"], "图2-3 系统模块主要逻辑关系图", 6.2)


def add_detailed_design(doc: Document, diagrams: dict[str, tuple[Path, Path]]) -> None:
    doc.add_heading("2.3 各个模块详细设计", level=2)
    fig_no = [4]
    groups = [
        ("2.3.1 登录注册模块", "登录注册模块负责所有用户进入系统前的身份校验和进入系统后的个人资料维护，是权限控制的入口。", "登录注册模块"),
        ("2.3.2 普通员工模块", "普通员工模块负责检测任务提交、结果查看和审核状态跟踪，是系统业务数据进入检测流程的源头。", "普通员工模块"),
        ("2.3.3 数据安全员模块", "数据安全员模块负责AI审核、规则维护、预警处理和风险等级配置，是系统安全治理能力最集中的模块。", "数据安全员模块"),
        ("2.3.4 系统管理员模块", "系统管理员模块负责用户、角色、权限、日志和配置维护，为系统运行和权限隔离提供基础支撑。", "系统管理员模块"),
    ]
    for heading, intro, module in groups:
        doc.add_heading(heading, level=3)
        base.add_para(doc, intro)
        for feature in [x for x in FEATURES if x["module"] == module]:
            add_function_detail(doc, feature, diagrams, fig_no)


def add_database_and_test(doc: Document, overview_diagrams: dict[str, Path]) -> None:
    doc.add_heading("2.4 数据库概念结构模型", level=2)
    base.add_para(doc, "数据库设计围绕用户权限、检测审核、安全治理和系统运维四类数据展开。用户权限类表负责登录和角色权限；检测审核类表负责员工提交内容、检测结果和主管审核；安全治理类表负责AI审核记录、敏感词、风险等级和预警；系统运维类表负责配置和日志。")
    doc.add_heading("2.4.1 表结构设计（数据库表结构表格）", level=3)
    tables = base.parse_schema()
    important = ["sys_user", "sys_role", "sys_permission", "detect_task", "detect_result", "audit_task", "audit_record", "warning_record", "ai_review_record", "sensitive_category", "sensitive_word", "risk_level_config", "system_config", "operation_log"]
    no = 1
    for table in important:
        if table not in tables:
            continue
        p = doc.add_paragraph(f"表2-{no} {table}表结构")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        base.add_table(doc, ["字段名", "类型", "是否可空", "字段说明"], [list(r) for r in tables[table]], widths=[3.5, 3.0, 2.0, 7.0])
        no += 1
    doc.add_heading("2.4.2 E-R模型", level=3)
    base.add_para(doc, "E-R模型描述系统核心实体及其关系。用户拥有角色，角色关联权限；用户提交检测任务，检测任务产生检测结果；检测结果进入审核任务并形成审核记录；高风险任务触发预警；安全员维护敏感词和风险等级，并保存AI审核记录；管理员查看操作日志和系统配置。")
    base.add_picture(doc, overview_diagrams["er"], "图2-40 E-R模型图", 6.2)

    doc.add_heading("3 系统实现与测试", level=1)
    base.add_para(doc, "系统采用前后端分离方式实现。前端Vue3项目负责页面交互和接口调用，后端Spring Boot项目负责认证、检测、审核、AI调用、规则维护和数据库访问。测试时分别使用employee、manager、security、admin四个账号验证不同角色功能。")
    test_rows = [
        ["登录与权限", "四个角色分别登录并访问其他角色页面", "只能访问自身角色菜单和接口", "通过"],
        ["员工文本检测", "提交包含手机号、邮箱、银行卡号的文本", "生成风险评分、命中摘要、脱敏结果和审核任务", "通过"],
        ["员工文件检测", "上传txt、pdf、doc或docx文件", "成功抽取文本并进入检测流程", "通过"],
        ["AI文本审核", "安全员提交业务文本", "返回结构化AI审核结论并保存审核记录", "通过"],
        ["AI文件审核", "安全员上传文件进行AI审核", "返回文件审核结论并保存文件名", "通过"],
        ["AI审核记录", "进入AI审核记录页面查看详情", "显示历史记录、原文和审核结论", "通过"],
        ["敏感词与风险等级", "新增敏感词和修改风险等级", "规则数据保存成功", "通过"],
        ["管理员运维", "查看用户、角色、权限、日志和配置", "管理页面正常展示", "通过"],
    ]
    base.add_table(doc, ["测试项", "测试步骤", "预期结果", "结论"], test_rows, widths=[3.0, 5.0, 6.0, 2.0])
    screenshots = [
        ("01-login.png", "图3-1 登录界面"),
        ("02-employee-detect.png", "图3-2 普通员工检测界面"),
        ("03-employee-history.png", "图3-3 普通员工历史记录界面"),
        ("09-security-ai.png", "图3-4 数据安全员AI审核界面"),
        ("10-security-sensitive.png", "图3-5 敏感词库管理界面"),
        ("11-security-warnings.png", "图3-6 风险预警管理界面"),
        ("12-security-risk-levels.png", "图3-7 风险等级配置界面"),
        ("13-admin-users.png", "图3-8 用户管理界面"),
        ("15-admin-permissions.png", "图3-9 权限管理界面"),
        ("16-admin-logs.png", "图3-10 日志管理界面"),
        ("17-admin-config.png", "图3-11 系统配置界面"),
    ]
    for file, cap in screenshots:
        base.add_picture(doc, SCREENSHOT_DIR / file, cap, 6.2)
    base.add_para(doc, "测试结果表明，系统能够完成从登录、权限隔离、员工检测、主管审核、安全员AI审核、AI审核记录查询、敏感词维护、风险等级维护到管理员运维的主要流程。新增AI审核记录后，安全员的审核动作具备了可追溯性，避免审核结果只停留在页面临时显示。")

    doc.add_heading("4 总结", level=1)
    for text in [
        "本次课程设计完成了智能数据守护者系统的需求分析、系统设计、详细设计、数据库建模、功能实现和测试验证。系统围绕企业数据安全场景，把敏感信息检测、AI审核、审核记录、风险预警、敏感词库、风险等级配置和系统运维整合到一个角色化平台中。",
        "新版详细设计按照功能逐项展开，每个功能都明确了参与者、输入输出、前端页面、接口地址、业务服务、涉及数据表、UML活动流程和业务模型。这样的设计方式能够更直接地说明功能如何从需求落到页面、接口、服务和数据库，也方便答辩时按模块逐一讲解。",
        "系统仍有继续优化空间，例如可以进一步增加部门主管模块的详细图示、完善真实分页查询、增强文件格式解析能力、增加AI审核结果结构化字段、补充更多自动化测试用例，并对角色权限做更细粒度的数据绑定。整体来看，系统已经具备课程设计要求的完整业务流程和工程实现基础。",
    ]:
        base.add_para(doc, text)


def count_cn(doc: Document) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", "\n".join(p.text for p in doc.paragraphs)))


def build() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    detailed = generate_detailed_diagrams()
    overview = base.generate_diagrams()
    doc = Document()
    base.configure_doc(doc)
    base.add_cover(doc)
    add_toc(doc)
    add_front_sections(doc, overview)
    add_detailed_design(doc, detailed)
    add_database_and_test(doc, overview)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("智能数据守护者系统课程设计报告")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    doc.save(OUT)
    print(OUT)
    print(count_cn(doc))
    print(len(doc.inline_shapes))
    print(len(doc.tables))


if __name__ == "__main__":
    build()
